import os
import re
import unicodedata
import tempfile
import logging
from datetime import datetime, timezone

from dotenv import load_dotenv
from jose import jwt as jose_jwt, JWTError
from supabase import create_client, Client as SupabaseClient

from fastapi import FastAPI, File, Form, UploadFile, HTTPException, BackgroundTasks, Depends, Security
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.staticfiles import StaticFiles

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
logger = logging.getLogger("giaoan")

# Import bảng tra cứu NLS chuẩn (109 mã từ bang-tra-cuu-nls-cua-hs.pdf)
from nls_database import (
    BANG_TRA_CUU, BANG_TRA_CUU_MAP,
    KEYWORD_TO_NLS_CODES, CATEGORY_NAMES,
    CODE_TO_TOOLS, ACTIVITY_TYPE_TOOLS,
    SUBJECT_NLS_PRIORITY,
)

app = FastAPI(title="Chuẩn Hóa Giáo Án")

# Serve thư mục static/ nếu có, hoặc thư mục gốc nếu không có
_app_dir    = os.path.dirname(os.path.abspath(__file__))
_static_dir = os.path.join(_app_dir, "static")
if os.path.isdir(_static_dir):
    app.mount("/static", StaticFiles(directory=_static_dir), name="static")

# Route riêng cho supabase.min.js (dù file ở root hay static/)
@app.get("/supabase.min.js", include_in_schema=False)
async def serve_supabase_js():
    for candidate in [
        os.path.join(_static_dir, "supabase.min.js"),   # static/supabase.min.js
        os.path.join(_app_dir,    "supabase.min.js"),   # root/supabase.min.js
    ]:
        if os.path.isfile(candidate):
            return FileResponse(candidate, media_type="application/javascript")
    # Fallback: redirect về CDN
    from fastapi.responses import RedirectResponse
    return RedirectResponse(
        "https://cdn.jsdelivr.net/npm/@supabase/supabase-js@2/dist/umd/supabase.min.js"
    )

# ═══════════════════════════════════════════════════════════════════════════════
# SUPABASE AUTH
# ═══════════════════════════════════════════════════════════════════════════════

load_dotenv()

SUPABASE_URL         = os.getenv("SUPABASE_URL", "")
SUPABASE_ANON_KEY    = os.getenv("SUPABASE_ANON_KEY", "")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY", "")
SUPABASE_JWT_SECRET  = os.getenv("SUPABASE_JWT_SECRET", "")
FREE_QUOTA           = int(os.getenv("FREE_QUOTA", "5"))

# AUTH_ENABLED = True khi có Supabase URL + anon key (không cần JWT_SECRET nữa)
AUTH_ENABLED: bool = bool(SUPABASE_URL and SUPABASE_ANON_KEY)

_supa: SupabaseClient | None = None


def _get_supa() -> SupabaseClient | None:
    """Lazy-init Supabase client dùng service key (bypass RLS)."""
    global _supa
    if _supa is None and SUPABASE_URL and SUPABASE_SERVICE_KEY:
        _supa = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
    return _supa


# ── Bearer token extractor ────────────────────────────────────────────────────
_bearer = HTTPBearer(auto_error=False)


async def get_current_user(
    creds: HTTPAuthorizationCredentials = Security(_bearer),
) -> dict:
    """
    Xác thực token do Supabase cấp.
    Ưu tiên dùng Supabase API (get_user) — không phụ thuộc JWT_SECRET.
    Fallback về jose_jwt nếu Supabase client chưa sẵn sàng.
    Khi AUTH_ENABLED = False → trả về user ảo cho dev local.
    """
    if not AUTH_ENABLED:
        return {"sub": "dev-local", "email": "dev@local.vn", "role": "authenticated"}

    if not creds:
        raise HTTPException(
            status_code=401,
            detail="Vui lòng đăng nhập để sử dụng tính năng này.",
        )

    token = creds.credentials

    # Phương án 1: Verify qua Supabase API (đáng tin cậy nhất)
    supa = _get_supa()
    if supa:
        try:
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
                future = ex.submit(supa.auth.get_user, token)
                resp = future.result(timeout=8)   # timeout 8 giây
            if resp and resp.user:
                return {
                    "sub":   resp.user.id,
                    "email": resp.user.email or "",
                }
        except concurrent.futures.TimeoutError:
            logger.warning("get_user timeout — thử JWT fallback")
        except Exception as e:
            logger.warning(f"get_user failed: {e}")

    # Phương án 2: Verify bằng JWT_SECRET (fallback)
    if SUPABASE_JWT_SECRET:
        try:
            payload = jose_jwt.decode(
                token,
                SUPABASE_JWT_SECRET,
                algorithms=["HS256"],
                audience="authenticated",
            )
            return payload
        except JWTError as e:
            logger.warning(f"JWT decode failed: {e}")

    raise HTTPException(
        status_code=401,
        detail="Phiên đăng nhập hết hạn. Vui lòng đăng nhập lại.",
    )


# ── Quota: kiểm tra + tăng trong một lần gọi duy nhất ───────────────────────

def check_and_increment_quota(user_id: str) -> None:
    """
    Kiểm tra quota và tăng used_quota ngay trong một transaction.

    Schema profiles:
      used_quota  INT     – số lần đã dùng trong tháng
      is_pro      BOOLEAN – True = không giới hạn
      quota_month TEXT    – 'YYYY-MM' của tháng hiện tại

    Logic:
      1. Lấy profile từ DB.
      2. Nếu is_pro → cho qua.
      3. Nếu quota_month ≠ tháng hiện tại → reset used_quota về 0.
      4. Nếu used_quota >= FREE_QUOTA → raise HTTP 402.
      5. Ngược lại → tăng used_quota + 1 và lưu lại.
    """
    if not AUTH_ENABLED or user_id == "dev-local":
        return

    supa = _get_supa()
    if not supa:
        return  # Không có Supabase → fail open

    current_month = datetime.now(timezone.utc).strftime("%Y-%m")

    try:
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
            future = ex.submit(
                lambda: supa.table("profiles")
                    .select("used_quota, is_pro, quota_month")
                    .eq("id", user_id)
                    .maybe_single()
                    .execute()
            )
            r = future.result(timeout=8)   # timeout 8 giây

        if not r.data:
            # Profile chưa có → tạo mới với used_quota = 1
            supa.table("profiles").upsert({
                "id": user_id,
                "used_quota":  1,
                "quota_month": current_month,
            }).execute()
            return

        data        = r.data
        is_pro      = data.get("is_pro", False)
        used        = data.get("used_quota", 0)
        quota_month = data.get("quota_month", "")

        # Pro → không giới hạn
        if is_pro:
            return

        # Sang tháng mới → reset bộ đếm
        if quota_month != current_month:
            used = 0

        # Kiểm tra giới hạn
        if used >= FREE_QUOTA:
            raise HTTPException(
                status_code=402,
                detail=(
                    f"Thầy/Cô đã dùng hết {FREE_QUOTA} lượt miễn phí tháng này. "
                    "Nâng cấp Pro để dùng không giới hạn."
                ),
            )

        # Tăng quota
        supa.table("profiles").update({
            "used_quota":  used + 1,
            "quota_month": current_month,
        }).eq("id", user_id).execute()

    except HTTPException:
        raise  # Truyền 402 lên
    except Exception as e:
        logger.warning(f"check_and_increment_quota: {e}")
        # Fail open: lỗi kết nối DB không chặn người dùng


# ── Namespaces ───────────────────────────────────────────────────────────────
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
O_NS    = "urn:schemas-microsoft-com:office:office"
V_NS    = "urn:schemas-microsoft-com:vml"

M_OMATHPARA = f"{{{MATH_NS}}}oMathPara"
M_OMATH     = f"{{{MATH_NS}}}oMath"
M_RUN       = f"{{{MATH_NS}}}r"
O_OLEOBJECT = f"{{{O_NS}}}OLEObject"
V_SHAPE     = f"{{{V_NS}}}shape"
W_OBJECT    = f"{{{W_NS}}}object"

MAX_CONTENT_EMU = int(Cm(16))

PRESERVED_STYLES = {
    "Normal", "Default Paragraph Font", "No Spacing",
    "Heading 1", "Heading 2", "Heading 3", "Heading 4",
    "Heading 5", "Heading 6", "Heading 7", "Heading 8", "Heading 9",
    "Title", "Subtitle", "Caption", "List Paragraph",
    "Table Normal", "Body Text", "Body Text 2", "Body Text 3",
    "Header", "Footer",
}
VML_UNIT_EMU = {"pt": 12700, "in": 914400, "cm": 360000, "mm": 36000, "px": 9525}


# ═══════════════════════════════════════════════════════════════════════════════
# DATABASE NĂNG LỰC SỐ — GDPT 2018
# ═══════════════════════════════════════════════════════════════════════════════

MON_LABELS = {
    "toan":      "Toán",
    "ngu_van":   "Ngữ văn",
    "tieng_anh": "Tiếng Anh",
    "vat_ly":    "Vật lý",
    "hoa_hoc":   "Hóa học",
    "sinh_hoc":  "Sinh học",
    "lich_su":   "Lịch sử",
    "dia_ly":    "Địa lý",
    "gdcd":      "GDCD / GDKTPL",
    "tin_hoc":   "Tin học",
    "cong_nghe": "Công nghệ",
    "khoa_hoc":  "Khoa học",
    "tnxh":      "Tự nhiên & Xã hội",
    "gdtc":      "GDTC",
    "am_nhac":   "Âm nhạc",
    "my_thuat":  "Mỹ thuật",
}

CAP_LABELS = {
    "tieu_hoc": "Tiểu học (Lớp 1–5)",
    "thcs":     "THCS (Lớp 6–9)",
    "thpt":     "THPT (Lớp 10–12)",
}

NL_SO_DB = {
    "toan": {
        "tieu_hoc": [
            "Thông tin & dữ liệu: Nhận biết và thu thập số liệu từ biểu đồ, bảng số trên học liệu số (sách giáo khoa điện tử, video học toán).",
            "Công cụ số: Dùng máy tính bỏ túi hoặc ứng dụng tính toán phù hợp lứa tuổi để kiểm tra kết quả bài toán.",
            "Quan sát số: Xem hình ảnh, video số về hình dạng và phép đo trong thực tế để liên hệ kiến thức toán học.",
            "An toàn số: Biết quy tắc sử dụng thiết bị số an toàn, có trách nhiệm trong học tập.",
        ],
        "thcs": [
            "Thông tin & dữ liệu: Tìm kiếm, đánh giá và sử dụng thông tin số đáng tin cậy để giải quyết bài toán thực tế.",
            "Công cụ toán học số: Khai thác phần mềm hình học động (GeoGebra, Desmos) để khám phá tính chất hình học và đồ thị hàm số.",
            "Xử lý dữ liệu: Dùng bảng tính (Excel/Google Sheets) để xử lý số liệu thống kê và vẽ biểu đồ nhận xét xu hướng.",
            "Giao tiếp & hợp tác số: Sử dụng nền tảng học trực tuyến để chia sẻ bài giải và học hợp tác nhóm.",
            "An toàn số: Bảo vệ thông tin cá nhân khi tham gia các nền tảng học toán trực tuyến.",
        ],
        "thpt": [
            "Khai thác dữ liệu: Thu thập dữ liệu thực tế từ nguồn số, làm sạch và phân tích bằng bảng tính hoặc Python/R.",
            "Công cụ toán học số: Sử dụng thành thạo GeoGebra, Wolfram Alpha, Desmos để khám phá, kiểm chứng và mô phỏng bài toán đại số, giải tích, xác suất.",
            "Lập trình ứng dụng toán: Viết chương trình (Python) để giải bài toán số, sắp xếp, tìm kiếm và mô phỏng mô hình toán học.",
            "Trình bày số: Tạo báo cáo, slide, poster kỹ thuật số trình bày kết quả nghiên cứu toán học rõ ràng và thuyết phục.",
            "Đạo đức số: Trích dẫn nguồn trung thực; nhận biết và phê phán thông tin thống kê sai lệch trên mạng.",
        ],
    },
    "ngu_van": {
        "tieu_hoc": [
            "Đọc số: Sử dụng sách điện tử, audio book, video đọc truyện để tiếp cận tác phẩm văn học phù hợp lứa tuổi.",
            "Tra cứu ngôn ngữ: Tra từ điển tiếng Việt trực tuyến để hiểu nghĩa từ và chính tả chính xác.",
            "Sáng tạo số: Bước đầu soạn thảo văn bản số đơn giản để chia sẻ suy nghĩ về tác phẩm.",
            "An toàn thông tin: Nhận biết nội dung không phù hợp lứa tuổi; báo cáo với người lớn khi gặp thông tin xấu trên mạng.",
        ],
        "thcs": [
            "Khai thác tài liệu số: Truy cập thư viện số, học liệu mở để đọc và nghiên cứu văn bản văn học.",
            "Soạn thảo & chỉnh sửa: Dùng Word/Google Docs để viết bài luận và định dạng văn bản đúng chuẩn.",
            "Đánh giá thông tin: Phân biệt nguồn đáng tin cậy khi nghiên cứu về tác giả, bối cảnh lịch sử và ý nghĩa tác phẩm.",
            "Tạo nội dung số: Thiết kế slide (PowerPoint/Canva), sơ đồ tư duy số để phân tích nhân vật, chủ đề, kết cấu tác phẩm.",
            "Giao tiếp văn học số: Tham gia diễn đàn văn học trực tuyến để chia sẻ cảm nhận và tranh luận có trách nhiệm.",
        ],
        "thpt": [
            "Nghiên cứu học thuật số: Khai thác cơ sở dữ liệu nghiên cứu (Google Scholar) để tìm tài liệu phê bình văn học; trích dẫn đúng chuẩn.",
            "Phân tích văn bản số: Sử dụng công cụ phân tích văn bản để khảo sát tần suất từ, văn phong và cấu trúc ngôn ngữ.",
            "Sáng tác đa phương tiện: Tạo sản phẩm số (podcast văn học, video book trailer, blog phê bình) thể hiện quan điểm cá nhân.",
            "Bản quyền & đạo đức số: Hiểu và thực hành nguyên tắc bản quyền khi sử dụng, trích dẫn và chia sẻ tác phẩm trong môi trường số.",
            "Phản biện thông tin: Nhận diện tin giả; xây dựng lập luận dựa trên bằng chứng số đáng tin cậy.",
        ],
    },
    "tieng_anh": {
        "tieu_hoc": [
            "Học liệu số: Dùng ứng dụng học tiếng Anh (Duolingo, BBC Kids) và video để luyện nghe và phát âm.",
            "Flashcard số: Dùng Quizlet để ghi nhớ từ vựng qua trò chơi tương tác phù hợp lứa tuổi.",
            "An toàn trực tuyến: Xin phép người lớn trước khi truy cập nội dung trực tuyến để học tiếng Anh.",
        ],
        "thcs": [
            "Môi trường số đa ngôn ngữ: Khai thác video, podcast (BBC Learning English, CNN 10) để luyện nghe – nói tiếng Anh trong bối cảnh thực.",
            "Công cụ ngôn ngữ số: Dùng từ điển số, công cụ dịch máy (DeepL) thông minh — so sánh, phân tích và không lệ thuộc hoàn toàn.",
            "Giao tiếp số đa văn hóa: Tham gia dự án trao đổi số (e-pals) để giao tiếp tiếng Anh với bạn quốc tế có trách nhiệm.",
            "Sáng tạo nội dung: Tạo video, audio clip, blog tiếng Anh để luyện kỹ năng viết – nói trong bối cảnh thực tế.",
        ],
        "thpt": [
            "Nghiên cứu học thuật: Đọc hiểu tài liệu tiếng Anh số (Wikipedia EN, Khan Academy) cho các dự án nghiên cứu.",
            "Công cụ AI ngôn ngữ: Dùng có phê phán Grammarly/ChatGPT để cải thiện bài viết — đánh giá độ chính xác, không sao chép.",
            "Truyền thông số đa ngôn ngữ: Tạo sản phẩm số tiếng Anh (vlog, podcast, newsletter) giới thiệu bản sắc văn hóa Việt Nam.",
            "Giao tiếp toàn cầu: Tham gia hội thảo trực tuyến quốc tế; hiểu phép lịch sự trong văn hóa số đa quốc gia.",
            "An toàn & bảo mật: Nhận biết lừa đảo trực tuyến tiếng Anh (phishing, scam); bảo vệ thông tin trên nền tảng quốc tế.",
        ],
    },
    "vat_ly": {
        "tieu_hoc": [],
        "thcs": [
            "Thí nghiệm ảo: Dùng PhET Interactive Simulations để khám phá hiện tượng vật lý an toàn và trực quan.",
            "Dữ liệu thực nghiệm: Dùng bảng tính để ghi, tính toán và vẽ đồ thị kết quả thí nghiệm; rút ra quy luật.",
            "Học liệu đa phương tiện: Khai thác video, animation vật lý để quan sát hiện tượng khó thực hiện tại lớp.",
            "Đánh giá thông tin: Phân biệt thông tin khoa học chính thống và thông tin phi khoa học trên mạng.",
        ],
        "thpt": [
            "Mô phỏng số: Dùng PhET, Algodoo và phần mềm mô phỏng vật lý để kiểm chứng định luật và mô hình hóa hiện tượng phức tạp.",
            "Xử lý dữ liệu thực nghiệm: Dùng Python/Excel xử lý số liệu, vẽ đồ thị, tính sai số và lập báo cáo khoa học.",
            "Tài nguyên học thuật số: Khai thác MIT OpenCourseWare, bài báo vật lý số để mở rộng kiến thức nâng cao.",
            "Thiết bị đo số: Dùng cảm biến số (data logger) thu thập dữ liệu chính xác và phân tích bằng phần mềm.",
            "Trình bày khoa học số: Tạo báo cáo, poster khoa học, video thuyết trình về kết quả thí nghiệm vật lý.",
        ],
    },
    "hoa_hoc": {
        "tieu_hoc": [],
        "thcs": [
            "Mô phỏng phản ứng: Dùng PhET, MolView để quan sát cấu trúc phân tử và mô phỏng phản ứng hóa học an toàn.",
            "Tra cứu hóa học: Khai thác PubChem, bảng tuần hoàn tương tác để tìm hiểu tính chất và ứng dụng nguyên tố.",
            "Thí nghiệm ảo: Dùng lab hóa học ảo (Labster) thực hành thí nghiệm nguy hiểm trong môi trường số.",
            "An toàn hóa chất: Đọc và hiểu bảng thông tin an toàn hóa chất số (SDS) để thực hành an toàn.",
        ],
        "thpt": [
            "Cơ sở dữ liệu hóa học: Khai thác PubChem, ChemSpider, NIST để tìm kiếm và phân tích tính chất, phổ của hợp chất.",
            "Mô phỏng phân tử 3D: Dùng Avogadro, ChemDraw để vẽ và phân tích cấu trúc phân tử và liên kết hóa học.",
            "Tính toán hóa học số: Dùng bảng tính để giải phương trình phản ứng, tính nồng độ, hiệu suất và phân tích số liệu.",
            "Tài nguyên học thuật: Đọc tóm tắt bài báo hóa học từ ACS Publications; đánh giá phương pháp và kết quả nghiên cứu.",
            "Trình bày hóa học số: Tạo báo cáo thực nghiệm số và video thuyết trình về quy trình tổng hợp hóa học.",
        ],
    },
    "sinh_hoc": {
        "tieu_hoc": [],
        "thcs": [
            "Tài liệu số: Khai thác video giải phẫu ảo, animation sinh học tế bào để quan sát cấu trúc và quá trình sinh học.",
            "Dữ liệu thực nghiệm: Dùng bảng tính ghi chép và phân tích số liệu thí nghiệm sinh học; tìm quy luật từ dữ liệu.",
            "Tra cứu khoa học: Tìm thông tin sinh học từ NCBI, Wikipedia khoa học; phân biệt thông tin chính xác và tin đồn.",
            "Giao tiếp khoa học số: Tạo poster số, infographic về chủ đề sinh học để chia sẻ có trách nhiệm.",
        ],
        "thpt": [
            "Cơ sở dữ liệu gene: Khai thác NCBI GenBank, Ensembl tra cứu trình tự gene và thông tin tiến hóa; hiểu cơ bản bioinformatics.",
            "Mô phỏng sinh học: Dùng NetLogo (mô hình quần thể), SnapGene (thao tác gene) để khám phá quá trình sinh học phức tạp.",
            "Phân tích dữ liệu: Dùng Python hoặc SPSS phân tích số liệu thực nghiệm sinh học và kiểm định giả thuyết.",
            "Tài nguyên học thuật: Đọc và phân tích tóm tắt bài báo từ PubMed, Nature; đánh giá bằng chứng khoa học.",
            "Truyền thông khoa học số: Tạo video, podcast khoa học về sức khỏe và môi trường — đảm bảo chính xác, không gây hoang mang.",
        ],
    },
    "lich_su": {
        "tieu_hoc": [
            "Tài liệu lịch sử số: Xem video, hình ảnh lịch sử trên kênh giáo dục uy tín để cảm nhận bối cảnh lịch sử sinh động.",
            "Bảo tàng ảo: Tham quan bảo tàng lịch sử ảo (Google Arts & Culture) để khám phá hiện vật và sự kiện lịch sử.",
            "Đánh giá thông tin: Bước đầu phân biệt thông tin lịch sử chính thống và thông tin không có cơ sở.",
        ],
        "thcs": [
            "Tư liệu số hóa: Khai thác ảnh tư liệu, bản đồ lịch sử, văn kiện số hóa từ kho lưu trữ để phân tích bối cảnh.",
            "Bảo tàng & di sản số: Dùng tour bảo tàng ảo (Google Arts & Culture, bảo tàng trực tuyến VN) để tìm hiểu di sản.",
            "Phân tích nguồn sử liệu: Đánh giá độ tin cậy thông tin lịch sử trên Internet; phân biệt nguồn sơ cấp, thứ cấp và nội dung sai lệch.",
            "Trình bày lịch sử số: Dùng timeline digital, sơ đồ tư duy, bản đồ số để tái hiện tiến trình và trình bày phân tích lịch sử.",
        ],
        "thpt": [
            "Nghiên cứu sử liệu số: Khai thác cơ sở dữ liệu lưu trữ số (National Archives, Wikisource, JSTOR) để tiếp cận tài liệu gốc.",
            "Bản đồ lịch sử số: Dùng GIS (ArcGIS Online, Google My Maps) để phân tích thay đổi lãnh thổ và di cư dân số.",
            "Đánh giá nguồn số: Phân tích tính xác thực và quan điểm của tư liệu số; nhận biết diễn giải lịch sử có định kiến.",
            "Truyền thông lịch sử: Tạo video tư liệu, podcast, blog về nhân vật và sự kiện lịch sử dựa trên bằng chứng đáng tin cậy.",
            "Di sản số: Tham gia số hóa di sản và thu thập ký ức cộng đồng qua tư liệu số địa phương.",
        ],
    },
    "dia_ly": {
        "tieu_hoc": [
            "Bản đồ số đơn giản: Xem bản đồ trực tuyến để nhận biết vị trí địa lý quê hương và đất nước.",
            "Hình ảnh địa lý số: Khai thác hình ảnh vệ tinh Google Earth để quan sát địa hình, sông núi thực tế.",
        ],
        "thcs": [
            "Bản đồ số: Dùng Google Maps, Google Earth để quan sát địa hình và phân bố không gian của các hiện tượng địa lý.",
            "Ảnh vệ tinh & khí hậu: Khai thác ảnh vệ tinh số và dữ liệu khí hậu trực tuyến (NASA Worldview) để phân tích hiện tượng địa lý.",
            "Thống kê địa lý: Dùng bảng tính xử lý dữ liệu dân số, kinh tế từ cơ sở dữ liệu số (World Bank, GSO Việt Nam).",
            "Trình bày địa lý số: Tạo bản đồ số, infographic và thuyết trình về phân tích địa lý một vùng lãnh thổ.",
        ],
        "thpt": [
            "GIS & bản đồ số: Dùng ArcGIS Online/QGIS để phân tích không gian, xây dựng bản đồ chuyên đề và diễn giải dữ liệu địa lý.",
            "Viễn thám: Khai thác và phân tích ảnh vệ tinh (Sentinel Hub, NASA Earthdata) để nghiên cứu biến đổi sử dụng đất và đô thị hóa.",
            "Cơ sở dữ liệu địa lý: Tra cứu và phân tích dữ liệu từ World Bank, UNDP, Tổng cục Thống kê để lập luận trong nghiên cứu.",
            "Mô hình hóa không gian: Dùng công cụ số mô phỏng tác động biến đổi khí hậu và thiên tai đến không gian địa lý.",
            "Truyền thông địa lý số: Tạo StoryMap, video địa lý và báo cáo đa phương tiện để chia sẻ nghiên cứu.",
        ],
    },
    "gdcd": {
        "tieu_hoc": [
            "Công dân số: Nhận biết quyền và trách nhiệm của trẻ em trong môi trường số; ứng xử lịch sự và an toàn.",
            "Thông tin an toàn: Phân biệt thông tin tốt và xấu; không chia sẻ thông tin cá nhân cho người lạ trực tuyến.",
            "Quy tắc mạng: Thực hành ứng xử văn minh, tôn trọng và thân thiện khi giao tiếp trực tuyến.",
        ],
        "thcs": [
            "Quyền & trách nhiệm số: Hiểu quyền công dân số, quyền bảo vệ dữ liệu cá nhân và trách nhiệm trong không gian mạng.",
            "Chống bạo lực mạng: Nhận biết, phòng tránh và ứng phó với cyberbullying; hỗ trợ bạn bè bị bắt nạt trực tuyến.",
            "Phân biệt tin giả: Phát triển tư duy phê phán để nhận diện tin giả và kiểm chứng thông tin trước khi chia sẻ.",
            "Đạo đức số: Tôn trọng bản quyền, quyền riêng tư và không phát tán thông tin có hại trên mạng.",
        ],
        "thpt": [
            "Pháp luật số: Hiểu Luật An ninh mạng, quyền bảo vệ dữ liệu cá nhân; áp dụng trong cuộc sống số hàng ngày.",
            "Tư duy phê phán số: Phân tích thông tin mạng xã hội; nhận diện chiến thuật thao túng, deepfake và diễn ngôn thù địch.",
            "Tham gia dân chủ số: Dùng nền tảng số để thảo luận chính sách và phản hồi xây dựng với cộng đồng có văn minh.",
            "Kinh tế số & quyền lao động: Nhận biết cơ hội và thách thức kinh tế số; hiểu quyền lao động trong môi trường làm việc số.",
            "An ninh thông tin cá nhân: Bảo vệ tài khoản số, nhận biết lừa đảo và quản lý danh tiếng số có trách nhiệm.",
        ],
    },
    "tin_hoc": {
        "tieu_hoc": [
            "Sử dụng thiết bị: Thao tác máy tính, máy tính bảng; biết bật/tắt, mở ứng dụng và lưu file đúng cách.",
            "Phần mềm cơ bản: Dùng phần mềm soạn thảo, vẽ và trình chiếu đơn giản để thể hiện ý tưởng học tập.",
            "Internet an toàn: Không chia sẻ mật khẩu và thông tin cá nhân; báo cáo nội dung không phù hợp.",
            "Tư duy tính toán: Nhận biết và giải bài toán đơn giản theo dạng thuật toán; lập trình khối kéo thả (Scratch).",
        ],
        "thcs": [
            "Kỹ năng văn phòng số: Sử dụng thành thạo bộ ứng dụng văn phòng (Microsoft 365/Google Workspace) cho học tập và dự án nhóm.",
            "Lập trình cơ bản: Viết chương trình Scratch và Python cơ bản để giải vấn đề có cấu trúc; hiểu vòng lặp, điều kiện và hàm.",
            "Tư duy tính toán: Phân tích vấn đề, tách bài toán lớn thành nhỏ; thiết kế thuật toán và trình bày bằng sơ đồ khối.",
            "Mạng & Internet: Hiểu nguyên lý Internet, mạng xã hội và dịch vụ số; sử dụng an toàn và có trách nhiệm.",
            "An ninh thông tin: Nhận biết virus, phishing, mật khẩu yếu; thực hành biện pháp bảo vệ thông tin cá nhân.",
        ],
        "thpt": [
            "Lập trình nâng cao: Dùng Python/JavaScript để giải bài toán phức tạp; làm quen lập trình hướng đối tượng.",
            "Khoa học dữ liệu: Thu thập, xử lý và trực quan hóa dữ liệu bằng Python (pandas, matplotlib) hoặc Power BI.",
            "AI & Machine Learning: Hiểu nguyên lý AI/ML; thực hành với Google Teachable Machine và các công cụ AI sẵn có.",
            "Phát triển phần mềm: Thiết kế và phát triển ứng dụng web/mobile đơn giản; dùng Git/GitHub kiểm soát phiên bản.",
            "An ninh mạng: Hiểu biện pháp bảo mật, nhận biết lỗ hổng cơ bản; thực hành đạo đức trong lĩnh vực bảo mật số.",
        ],
    },
    "cong_nghe": {
        "tieu_hoc": [
            "Thiết bị công nghệ: Nhận biết và sử dụng an toàn thiết bị công nghệ trong gia đình và trường học.",
            "Sản phẩm số đơn giản: Tạo thiệp điện tử, tranh số bằng phần mềm vẽ phù hợp lứa tuổi.",
            "An toàn công nghệ: Nhận biết nguy hiểm điện từ và thực hành quy tắc an toàn khi sử dụng thiết bị.",
        ],
        "thcs": [
            "Thiết kế số: Dùng phần mềm CAD 2D/3D đơn giản (Tinkercad, SketchUp Free) để thiết kế và mô phỏng sản phẩm kỹ thuật.",
            "Lập trình ứng dụng: Lập trình vi điều khiển (Arduino, micro:bit) để điều khiển thiết bị và tự động hóa đơn giản.",
            "Internet of Things: Tìm hiểu và thực hành dự án IoT cơ bản; kết nối thiết bị với Internet để thu thập dữ liệu.",
            "Công nghệ số trong sản xuất: Tìm hiểu ứng dụng in 3D, robot, CNC trong sản xuất hiện đại; nhận xét tác động đến lao động.",
        ],
        "thpt": [
            "Thiết kế kỹ thuật số: Dùng Fusion 360/SolidWorks để thiết kế, mô phỏng và tối ưu sản phẩm kỹ thuật chuyên nghiệp.",
            "Hệ thống nhúng & IoT: Lập trình Raspberry Pi/Arduino để thiết kế dự án IoT hoàn chỉnh từ thu thập đến điều khiển.",
            "Tự động hóa & Robot: Thiết kế và lập trình robot thực hiện nhiệm vụ cụ thể; hiểu nguyên lý điều khiển tự động.",
            "Sản xuất thông minh: Tìm hiểu Công nghiệp 4.0 và nhà máy thông minh; đánh giá tác động kinh tế-xã hội.",
            "Khởi nghiệp công nghệ: Xây dựng prototype sản phẩm công nghệ giải quyết vấn đề thực tế; trình bày ý tưởng kinh doanh số.",
        ],
    },
    "khoa_hoc": {
        "tieu_hoc": [
            "Quan sát số: Dùng video, hình ảnh số về thiên nhiên (thời tiết, động thực vật) bổ sung quan sát trực tiếp.",
            "Ghi chép số: Chụp ảnh, quay video thí nghiệm đơn giản để lưu lại quá trình quan sát khoa học.",
            "Tra cứu khoa học: Dùng công cụ tìm kiếm để tìm hiểu hiện tượng tự nhiên; phân biệt trang web đáng tin cậy dành cho trẻ em.",
            "An toàn số: Luôn có người lớn hướng dẫn khi dùng Internet để học khoa học; nhận biết nội dung không phù hợp.",
        ],
        "thcs": [],
        "thpt": [],
    },
    "tnxh": {
        "tieu_hoc": [
            "Khám phá số: Dùng video, hình ảnh số để tìm hiểu về con người, thiên nhiên và xã hội ngoài lớp học.",
            "Bản đồ số đơn giản: Xem Google Maps để nhận biết vị trí địa lý địa phương, đất nước và thế giới.",
            "Giao tiếp an toàn: Ứng xử lịch sự trong giao tiếp trực tuyến; nhận biết nguy cơ từ người lạ qua mạng.",
            "Sức khỏe số: Thực hành thời gian sử dụng màn hình hợp lý; biết tác hại của dùng thiết bị quá nhiều với sức khỏe.",
        ],
        "thcs": [],
        "thpt": [],
    },
    "gdtc": {
        "tieu_hoc": [
            "Theo dõi sức khỏe số: Dùng ứng dụng đơn giản (đếm bước chân) để theo dõi hoạt động thể chất hàng ngày.",
            "Video thể dục: Xem video hướng dẫn bài tập từ nguồn đáng tin cậy để luyện tập đúng kỹ thuật tại nhà.",
        ],
        "thcs": [
            "Theo dõi thể lực số: Dùng ứng dụng và thiết bị đeo thông minh để theo dõi, ghi chép và phân tích chỉ số thể lực.",
            "Phân tích kỹ thuật số: Quay video slow-motion phân tích kỹ thuật động tác thể thao để cải thiện kỹ năng.",
            "Thông tin dinh dưỡng: Tra cứu và đánh giá thông tin dinh dưỡng từ nguồn y tế uy tín; phân biệt quảng cáo sai sự thật.",
        ],
        "thpt": [
            "Phân tích hiệu suất số: Dùng thiết bị thể thao thông minh để theo dõi, phân tích và lập kế hoạch tập luyện dựa trên dữ liệu.",
            "Truyền thông thể thao số: Tạo nội dung số (video hướng dẫn, blog sức khỏe) chia sẻ lối sống lành mạnh có trách nhiệm.",
            "Đánh giá thông tin sức khỏe: Đánh giá phê phán thông tin dinh dưỡng, thực phẩm chức năng trên mạng; nhận biết quảng cáo phóng đại.",
        ],
    },
    "am_nhac": {
        "tieu_hoc": [
            "Nghe nhạc số: Khai thác nền tảng nhạc số phù hợp lứa tuổi để nghe và nhận biết các thể loại âm nhạc Việt Nam và quốc tế.",
            "Nhạc cụ ảo: Trải nghiệm nhạc cụ số (piano ảo, gõ nhịp) qua ứng dụng để cảm nhận âm thanh và nhịp điệu.",
            "Video âm nhạc: Xem video biểu diễn và học nhạc từ nguồn đáng tin cậy để quan sát kỹ thuật.",
        ],
        "thcs": [
            "Sáng tác số: Dùng GarageBand, BandLab để thử nghiệm sáng tác và phối khí điện tử.",
            "Phân tích âm nhạc số: Dùng công cụ phân tích âm thanh số để quan sát sóng âm, nhận biết nhịp, cao độ và cấu trúc.",
            "Bản quyền âm nhạc: Hiểu bản quyền âm nhạc và Creative Commons; sử dụng và chia sẻ nhạc số hợp pháp.",
        ],
        "thpt": [
            "Sản xuất âm nhạc số: Dùng DAW (GarageBand, FL Studio, Ableton) để sáng tác, thu âm, mix và master sản phẩm cá nhân.",
            "Phân tích học thuật: Khai thác tài liệu số về lý thuyết âm nhạc để phân tích phong cách và kỹ thuật nhạc sĩ.",
            "Phân phối số: Hiểu hệ sinh thái âm nhạc số (streaming, YouTube); biết cách phân phối sản phẩm hợp pháp và hiệu quả.",
        ],
    },
    "my_thuat": {
        "tieu_hoc": [
            "Vẽ số: Thử nghiệm vẽ và tô màu bằng phần mềm vẽ đơn giản (Paint, Sketchpad) để thể hiện ý tưởng sáng tạo.",
            "Tham quan nghệ thuật số: Khám phá bảo tàng nghệ thuật ảo (Google Arts & Culture) để chiêm ngưỡng tác phẩm nổi tiếng.",
            "Sáng tạo số đơn giản: Tạo thiệp, tranh số đơn giản chia sẻ với gia đình và bạn bè.",
        ],
        "thcs": [
            "Vẽ kỹ thuật số: Dùng Krita, Procreate, Canva để tạo tác phẩm nghệ thuật số với kỹ thuật layering.",
            "Nhiếp ảnh & chỉnh sửa: Thực hành chụp ảnh sáng tạo và chỉnh sửa cơ bản (cắt xén, màu sắc, ánh sáng).",
            "Bản quyền hình ảnh: Hiểu và tôn trọng bản quyền; dùng hình ảnh Creative Commons và ghi nguồn trung thực.",
            "Portfolio số: Tạo portfolio nghệ thuật số để trưng bày và chia sẻ tác phẩm có trách nhiệm.",
        ],
        "thpt": [
            "Đồ họa chuyên nghiệp: Dùng Adobe Photoshop/Illustrator hoặc GIMP/Inkscape để tạo tác phẩm đồ họa chuyên nghiệp.",
            "Nghệ thuật kỹ thuật số: Phát triển phong cách nghệ thuật số cá nhân; thử nghiệm generative art và AI art với tư duy phê phán.",
            "Portfolio chuyên nghiệp: Xây dựng portfolio trực tuyến để trình bày quá trình sáng tác và tác phẩm.",
            "Thẩm mỹ & đạo đức số: Phân tích tác động thẩm mỹ – xã hội của hình ảnh số; thực hành đạo đức khi chia sẻ hình ảnh.",
        ],
    },
}


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _apply_font(run, size_pt: float = 14.0):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), "Times New Roman")


def _set_spacing(para):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before      = Pt(3)
    pf.space_after       = Pt(3)


def _center_wp(wp_elem):
    pPr = wp_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        wp_elem.insert(0, pPr)
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "center")


def _ancestor_wp(element):
    node = element.getparent()
    while node is not None:
        if node.tag == f"{{{W_NS}}}p":
            return node
        node = node.getparent()
    return None


# ═══════════════════════════════════════════════════════════════════════════════
# 1–10: CÁC HÀM CHUẨN HÓA (giữ nguyên như cũ)
# ═══════════════════════════════════════════════════════════════════════════════

def set_page_margins(doc):
    for section in doc.sections:
        section.page_height = Cm(29.7); section.page_width  = Cm(21.0)
        section.top_margin  = Cm(2.0);  section.bottom_margin= Cm(2.0)
        section.left_margin = Cm(3.0);  section.right_margin = Cm(2.0)


def _clean_one_para(para, doc):
    style_name = para.style.name if para.style else "Normal"
    if style_name not in PRESERVED_STYLES:
        try: para.style = doc.styles["Normal"]
        except Exception: pass
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        shd = pPr.find(qn("w:shd"))
        if shd is not None: pPr.remove(shd)
    for run in para.runs:
        rPr = run._r.find(qn("w:rPr"))
        if rPr is not None:
            for bad in (qn("w:shd"), qn("w:highlight")):
                el = rPr.find(bad)
                if el is not None: rPr.remove(el)


def clean_paragraph_styles(doc):
    for para in doc.paragraphs: _clean_one_para(para, doc)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs: _clean_one_para(para, doc)


def format_paragraph(para):
    _set_spacing(para)
    for run in para.runs: _apply_font(run, 14.0)


def set_table_autofit(doc):
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None: tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
        tblW.set(qn("w:w"), "5000"); tblW.set(qn("w:type"), "pct")
        tblL = tblPr.find(qn("w:tblLayout"))
        if tblL is None: tblL = OxmlElement("w:tblLayout"); tblPr.append(tblL)
        tblL.set(qn("w:type"), "autofit")
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs: format_paragraph(para)


def fix_math_to_inline(doc):
    body = doc.element.body
    for wp in body.iter(f"{{{W_NS}}}p"):
        for child in list(wp):
            if child.tag == M_OMATHPARA:
                idx = list(wp).index(child)
                for oMath in child.findall(M_OMATH):
                    child.remove(oMath); wp.insert(idx, oMath); idx += 1
                wp.remove(child)


def format_math_runs(doc):
    body = doc.element.body
    for oMath in body.iter(M_OMATH):
        for mRun in oMath.iter(M_RUN):
            w_rPr = mRun.find(qn("w:rPr"))
            if w_rPr is None: w_rPr = OxmlElement("w:rPr"); mRun.insert(0, w_rPr)
            rFonts = w_rPr.find(qn("w:rFonts"))
            if rFonts is None: rFonts = OxmlElement("w:rFonts"); w_rPr.insert(0, rFonts)
            for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rFonts.set(qn(attr), "Times New Roman")
            # Dùng string "w:sz" cho OxmlElement, qn("w:sz") cho find — không nhầm lẫn
            for tag_str in ("w:sz", "w:szCs"):
                el = w_rPr.find(qn(tag_str))
                if el is None: el = OxmlElement(tag_str); w_rPr.append(el)
                el.set(qn("w:val"), "28")


def resize_inline_images(doc):
    for shape in doc.inline_shapes:
        try:
            if shape.width > MAX_CONTENT_EMU:
                ratio = MAX_CONTENT_EMU / shape.width
                shape.width = MAX_CONTENT_EMU; shape.height = int(shape.height * ratio)
            wp = _ancestor_wp(shape._inline)
            if wp is not None: _center_wp(wp)
        except Exception as e: logger.warning(f"resize_inline_images: {e}")


def format_header_footer(doc):
    _GETTERS = [lambda s: s.header, lambda s: s.footer,
                lambda s: s.even_page_header, lambda s: s.even_page_footer,
                lambda s: s.first_page_header, lambda s: s.first_page_footer]
    for section in doc.sections:
        for getter in _GETTERS:
            try:
                hf = getter(section)
                if hf:
                    for para in hf.paragraphs:
                        for run in para.runs: _apply_font(run, 13.0)
            except Exception: pass


def _wincom_convert_mathtype(input_path, output_path):
    try:
        import win32com.client as win32, pythoncom
    except ImportError:
        return False
    word = None
    try:
        pythoncom.CoInitialize()
        word = win32.Dispatch("Word.Application")
        word.Visible = False; word.DisplayAlerts = 0
        wdoc = word.Documents.Open(os.path.abspath(input_path),
                                   ReadOnly=False, AddToRecentFiles=False,
                                   ConfirmConversions=False)
        MAX_W = (16.0 / 2.54) * 72
        for shape in list(wdoc.InlineShapes):
            try:
                ole = shape.OLEFormat; prog = str(ole.ProgID) if ole else ""
                if "Equation" not in prog: continue
                converted = False
                for cid in ("MathType.MathML.7","MathType.MathML.6","MathType.MathML.5"):
                    try: ole.ConvertTo(ClassType=cid, DisplayAsIcon=False); converted=True; break
                    except: pass
                if not converted and shape.Width > MAX_W:
                    r = MAX_W/shape.Width; shape.Height=int(shape.Height*r); shape.Width=int(MAX_W)
            except: pass
        wdoc.SaveAs2(os.path.abspath(output_path), FileFormat=12)
        wdoc.Close(SaveChanges=False); word.Quit(); return True
    except Exception as e:
        logger.warning(f"win32com: {e}")
        try: word.Quit()
        except: pass
        return False
    finally:
        try:
            import pythoncom; pythoncom.CoUninitialize()
        except: pass


def fix_mathtype_ole_fallback(doc):
    body = doc.element.body
    for w_obj in body.iter(W_OBJECT):
        if not any("Equation" in ole.get("ProgID","") for ole in w_obj.iter(O_OLEOBJECT)):
            continue
        for vshape in w_obj.iter(V_SHAPE):
            style = vshape.get("style","")
            if not style: continue
            wm = re.search(r"width\s*:\s*([\d.]+)\s*(pt|in|cm|mm|px)", style, re.I)
            hm = re.search(r"height\s*:\s*([\d.]+)\s*(pt|in|cm|mm|px)", style, re.I)
            if not wm: continue
            wv,wu = float(wm.group(1)), wm.group(2).lower()
            if wv * VML_UNIT_EMU.get(wu,12700) > MAX_CONTENT_EMU:
                ratio = MAX_CONTENT_EMU / (wv * VML_UNIT_EMU[wu])
                ns = re.sub(r"width\s*:\s*[\d.]+\s*(?:pt|in|cm|mm|px)",
                            f"width:{wv*ratio:.3f}{wu}", style, flags=re.I)
                if hm: ns = re.sub(r"height\s*:\s*[\d.]+\s*(?:pt|in|cm|mm|px)",
                                   f"height:{float(hm.group(1))*ratio:.3f}{hm.group(2)}",
                                   ns, flags=re.I)
                vshape.set("style", ns)
        wr = w_obj.getparent()
        if wr is not None:
            wp = wr.getparent()
            if wp is not None and wp.tag == f"{{{W_NS}}}p": _center_wp(wp)


def _auto_insert_nls_in_activities(doc: Document, mon: str = "") -> None:
    """
    TỰ ĐỘNG gọi inject_competence_to_docx với selected_items rỗng.
    Hàm sẽ tự phát hiện đoạn dùng công cụ số và gợi ý mã NLS inline.
    Dùng cho Tab 1 (Chuẩn hóa ngay).

    Bỏ qua nếu bảng đã có hàng 'Năng lực số tích hợp' (tránh trùng).
    """
    # Kiểm tra đã chèn NLS vào bảng chưa
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if "năng lực số tích hợp" in cell.text.lower():
                    logger.info("Bảng đã có NLS — bỏ qua auto-insert.")
                    return

    # Phát hiện môn từ nội dung file (nếu chưa biết)
    detected_mon = mon or _detect_subject_from_doc(doc)

    # Trích xuất từng hoạt động + gợi ý mã NLS theo nội dung + môn học
    act_sections = _extract_activities_with_nls(doc, mon=detected_mon)
    if not act_sections:
        logger.info("Không phát hiện hoạt động nào trong giáo án.")
        return

    # Chuyển sang định dạng selected_items để tái dùng insert_nls_into_activity_tables
    selected_items: list[dict] = []
    for act in act_sections:
        for code_item in act.get("codes", []):
            selected_items.append({
                "activity_name": act["name"],
                "activity_type": act.get("type", "other"),  # quan trọng: loại HĐ
                "text":          code_item.get("text",  ""),
                "code":          code_item.get("code",  ""),
                "tools":         code_item.get("tools", []),
            })

    if not selected_items:
        return

    # Lấy tên môn hiển thị để truyền vào nhãn NLS (biến động)
    subject_label = MON_LABELS.get(detected_mon, "")

    logger.info(f"Auto-insert NLS: {len(act_sections)} HĐ, {len(selected_items)} mã, môn={subject_label or '?'}")

    # Gọi hàm chèn NLS với tên môn biến động
    inject_competence_to_docx(
        doc, selected_items,
        mon_label=subject_label,
        cap_label="",
    )


def process_docx(input_path, output_path):
    import traceback
    doc = Document(input_path)

    # Bước 0: $...$ → OMML thật (từ Gemini/ChatGPT)
    try:
        convert_latex_in_doc(doc)
    except Exception as e:
        logger.error(f"convert_latex_in_doc error:\n{traceback.format_exc()}")

    clean_paragraph_styles(doc)
    set_page_margins(doc)
    fix_math_to_inline(doc)
    for para in doc.paragraphs: format_paragraph(para)
    set_table_autofit(doc)
    format_math_runs(doc)
    resize_inline_images(doc)
    format_header_footer(doc)
    fix_mathtype_ole_fallback(doc)

    # Tự động chèn NLS phù hợp bối cảnh vào bảng tiến trình từng hoạt động
    try:
        _auto_insert_nls_in_activities(doc)
    except Exception as e:
        logger.warning(f"auto_insert_nls: {e}\n{traceback.format_exc()}")

    doc.save(output_path)


# ═══════════════════════════════════════════════════════════════════════════════
# NĂNG LỰC SỐ — INSERTION LOGIC
# ═══════════════════════════════════════════════════════════════════════════════

def _make_wp(text: str, bold: bool = False, indent_twips: int = 0) -> object:
    """Tạo một <w:p> XML element chuẩn Times New Roman 14pt."""
    p = OxmlElement("w:p")

    pPr = OxmlElement("w:pPr")
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), "60"); spc.set(qn("w:after"), "60")
    spc.set(qn("w:line"), "240"); spc.set(qn("w:lineRule"), "auto")  # Single
    pPr.append(spc)
    if indent_twips:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent_twips))
        pPr.append(ind)
    p.append(pPr)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), "Times New Roman")
    rPr.append(rFonts)
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag); el.set(qn("w:val"), "28"); rPr.append(el)
    if bold:
        rPr.append(OxmlElement("w:b"))
    r.append(rPr)

    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t); p.append(r)
    return p


def _find_insertion_wp(doc: Document):
    """
    Tìm đoạn văn CUỐI của phần 'Năng lực' để chèn NLS vào SAU ĐÓ.

    Thuật toán:
    1. Tìm vị trí bắt đầu của phần 'Năng lực' (heading hoặc dòng chứa 'năng lực').
    2. Quét tiếp đến khi gặp dòng đánh dấu kết thúc phần đó
       (ví dụ: 'phẩm chất', 'thái độ', 'thiết bị', 'chuẩn bị', 'tiến trình').
    3. Đoạn ngay trước khi gặp keyword kết thúc chính là vị trí chèn.

    Kết quả: NLS được chèn vào CUỐI phần Năng lực, không phải ở giữa.
    """
    paras = doc.paragraphs

    # ── Từ khóa kết thúc phần Năng lực ────────────────────────────────────────
    # Khi gặp những từ này → phần Năng lực đã kết thúc
    END_SECTION_KW = [
        "phẩm chất", "thái độ", "đức tính",
        "thiết bị", "học liệu", "đồ dùng",
        "chuẩn bị", "phương tiện",
        "tiến trình", "tổ chức hoạt động",
        "phương pháp", "hình thức",
        "kiểm tra", "đánh giá",
    ]

    # ── Bước 1: Tìm vị trí đoạn bắt đầu phần 'Năng lực' ──────────────────────
    nl_start_idx: int | None = None
    for i, para in enumerate(paras[:60]):
        tl = para.text.strip().lower()
        if "năng lực" in tl:
            nl_start_idx = i   # ghi nhận CUỐI cùng trong vùng tìm kiếm

    if nl_start_idx is None:
        # Fallback 1: chèn sau đoạn chứa 'mục tiêu'
        for para in paras[:40]:
            if "mục tiêu" in para.text.lower():
                return para._p
        # Fallback 2: đoạn không rỗng đầu tiên
        for para in paras:
            if para.text.strip():
                return para._p
        return None

    # ── Bước 2: Quét tiếp từ nl_start_idx để tìm cuối phần Năng lực ──────────
    last_in_section = nl_start_idx
    for i in range(nl_start_idx + 1, min(nl_start_idx + 60, len(paras))):
        tl = paras[i].text.strip().lower()
        if not tl:
            continue  # bỏ qua dòng trống

        # Phát hiện keyword kết thúc phần Năng lực
        if any(kw in tl for kw in END_SECTION_KW):
            break

        last_in_section = i   # vẫn trong phần Năng lực → cập nhật

    logger.debug(f"_find_insertion_wp: chèn sau đoạn [{last_in_section}] "
                 f"'{paras[last_in_section].text.strip()[:60]}'")
    return paras[last_in_section]._p


def insert_nang_luc_so(doc: Document, competencies: list, mon_label: str, cap_label: str):
    """
    Chèn khối Năng lực số vào file Word sau điểm tìm được.
    Bỏ qua nếu file đã có 'năng lực số' (tránh chèn trùng).
    """
    # Kiểm tra xem đã có chưa
    for para in doc.paragraphs:
        if "năng lực số" in para.text.lower():
            logger.info("Đã tồn tại mục 'Năng lực số' — bỏ qua chèn trùng.")
            return

    ref_p = _find_insertion_wp(doc)
    if ref_p is None:
        logger.warning("Không tìm thấy vị trí chèn phù hợp.")
        return

    # Tạo các đoạn (theo thứ tự ngược vì dùng addnext liên tiếp)
    header  = _make_wp(f"- Năng lực số tích hợp ({mon_label} – {cap_label}):", bold=True)
    items   = [_make_wp(f"+ {c}", indent_twips=360) for c in competencies]

    # Chèn: header trước, rồi các mục theo thứ tự
    current = ref_p
    current.addnext(header); current = header
    for item in items:
        current.addnext(item); current = item


# ═══════════════════════════════════════════════════════════════════════════════
# LATEX → OMML CONVERTER  (xử lý $...$ trong file từ AI như Gemini/ChatGPT)
# ═══════════════════════════════════════════════════════════════════════════════

MML_NS_URI  = "http://www.w3.org/1998/Math/MathML"
OMML_NS_URI = MATH_NS          # đã định nghĩa ở trên

def _mq(tag: str) -> str: return f"{{{OMML_NS_URI}}}{tag}"

def _omr(text: str):
    """Tạo m:r (math run) chứa text."""
    r = etree.Element(_mq("r"))
    t = etree.SubElement(r, _mq("t"))
    t.text = text
    return r

def _mml_tag(elem) -> str:
    tag = elem.tag
    return tag.split("}", 1)[1] if isinstance(tag, str) and "}" in tag else tag

def _mml_kids(mml_elem) -> list:
    result = []
    for child in mml_elem:
        result.extend(_mml_to_omml(child))
    return result

def _mml_to_omml(mml_elem) -> list:
    """
    Chuyển đổi đệ quy MathML → OMML.
    Xử lý các cấu trúc phổ biến trong toán học phổ thông VN.
    """
    tag = _mml_tag(mml_elem)

    # ── Containers → xử lý children ─────────────────────────────────────────
    if tag in ("math", "mrow", "mstyle", "mpadded", "semantics",
               "annotation-xml", "mphantom"):
        return _mml_kids(mml_elem)

    # ── Phân số: \frac → m:f ─────────────────────────────────────────────────
    if tag == "mfrac":
        ch = list(mml_elem)
        f   = etree.Element(_mq("f"))
        num = etree.SubElement(f, _mq("num"))
        den = etree.SubElement(f, _mq("den"))
        if ch:     [num.append(c) for c in _mml_to_omml(ch[0])]
        if len(ch) > 1: [den.append(c) for c in _mml_to_omml(ch[1])]
        return [f]

    # ── Lũy thừa: x^n → m:sSup ──────────────────────────────────────────────
    if tag == "msup":
        ch = list(mml_elem)
        s   = etree.Element(_mq("sSup"))
        e   = etree.SubElement(s, _mq("e"))
        sup = etree.SubElement(s, _mq("sup"))
        if ch:     [e.append(c)   for c in _mml_to_omml(ch[0])]
        if len(ch) > 1: [sup.append(c) for c in _mml_to_omml(ch[1])]
        return [s]

    # ── Chỉ số dưới: x_n → m:sSub ───────────────────────────────────────────
    if tag == "msub":
        ch = list(mml_elem)
        s   = etree.Element(_mq("sSub"))
        e   = etree.SubElement(s, _mq("e"))
        sub = etree.SubElement(s, _mq("sub"))
        if ch:     [e.append(c)   for c in _mml_to_omml(ch[0])]
        if len(ch) > 1: [sub.append(c) for c in _mml_to_omml(ch[1])]
        return [s]

    # ── Lũy thừa + chỉ số: x_n^m → m:sSubSup ───────────────────────────────
    if tag == "msubsup":
        ch  = list(mml_elem)
        s   = etree.Element(_mq("sSubSup"))
        e   = etree.SubElement(s, _mq("e"))
        sub = etree.SubElement(s, _mq("sub"))
        sup = etree.SubElement(s, _mq("sup"))
        if ch:     [e.append(c)   for c in _mml_to_omml(ch[0])]
        if len(ch) > 1: [sub.append(c) for c in _mml_to_omml(ch[1])]
        if len(ch) > 2: [sup.append(c) for c in _mml_to_omml(ch[2])]
        return [s]

    # ── Căn bậc 2: \sqrt → m:rad ─────────────────────────────────────────────
    if tag == "msqrt":
        rad  = etree.Element(_mq("rad"))
        pr   = etree.SubElement(rad, _mq("radPr"))
        dh   = etree.SubElement(pr,  _mq("degHide")); dh.set(_mq("val"), "1")
        etree.SubElement(rad, _mq("deg"))             # bậc rỗng
        e    = etree.SubElement(rad, _mq("e"))
        [e.append(c) for c in _mml_kids(mml_elem)]
        return [rad]

    # ── Căn bậc n: \sqrt[n]{x} → m:rad ──────────────────────────────────────
    if tag == "mroot":
        ch  = list(mml_elem)
        rad = etree.Element(_mq("rad"))
        deg = etree.SubElement(rad, _mq("deg"))
        e   = etree.SubElement(rad, _mq("e"))
        if len(ch) > 1: [deg.append(c) for c in _mml_to_omml(ch[1])]  # bậc
        if ch:          [e.append(c)   for c in _mml_to_omml(ch[0])]  # số bị khai căn
        return [rad]

    # ── Dấu ngoặc: \left( \right) → m:d ────────────────────────────────────
    if tag == "mfenced":
        open_ch  = mml_elem.get("open",  "(")
        close_ch = mml_elem.get("close", ")")
        d    = etree.Element(_mq("d"))
        dPr  = etree.SubElement(d, _mq("dPr"))
        bc   = etree.SubElement(dPr, _mq("begChr")); bc.set(_mq("val"), open_ch)
        ec   = etree.SubElement(dPr, _mq("endChr")); ec.set(_mq("val"), close_ch)
        for child in mml_elem:
            e = etree.SubElement(d, _mq("e"))
            [e.append(c) for c in _mml_to_omml(child)]
        return [d]

    # ── Giới hạn trên: \overset → m:limUpp ──────────────────────────────────
    if tag == "mover":
        ch  = list(mml_elem)
        lim = etree.Element(_mq("limUpp"))
        e   = etree.SubElement(lim, _mq("e"))
        lv  = etree.SubElement(lim, _mq("lim"))
        if ch:     [e.append(c)  for c in _mml_to_omml(ch[0])]
        if len(ch) > 1: [lv.append(c) for c in _mml_to_omml(ch[1])]
        return [lim]

    # ── Giới hạn dưới → m:limLow ─────────────────────────────────────────────
    if tag == "munder":
        ch  = list(mml_elem)
        lim = etree.Element(_mq("limLow"))
        e   = etree.SubElement(lim, _mq("e"))
        lv  = etree.SubElement(lim, _mq("lim"))
        if ch:     [e.append(c)  for c in _mml_to_omml(ch[0])]
        if len(ch) > 1: [lv.append(c) for c in _mml_to_omml(ch[1])]
        return [lim]

    # ── Text/ký tự/số → m:r ──────────────────────────────────────────────────
    if tag in ("mi", "mn", "mo", "mtext", "ms"):
        txt = (mml_elem.text or "").strip()
        return [_omr(txt)] if txt else []

    # ── Mặc định: xử lý children ─────────────────────────────────────────────
    return _mml_kids(mml_elem)


def _latex_to_omath(latex: str):
    """
    Chuyển chuỗi LaTeX thành phần tử m:oMath (OMML).
    Trả về None nếu chuyển đổi thất bại.
    """
    try:
        from latex2mathml.converter import convert as l2m
        mathml_str = l2m(latex)
        mml_elem   = etree.fromstring(mathml_str.encode("utf-8"))
        omath      = etree.Element(f"{{{MATH_NS}}}oMath")
        [omath.append(c) for c in _mml_to_omml(mml_elem)]
        return omath
    except Exception as e:
        logger.warning(f"LaTeX→OMML thất bại cho '{latex}': {e}")
        return None


_LATEX_PAT = re.compile(r"\$\$(.+?)\$\$|\$([^$\n]+?)\$", re.DOTALL)

def _process_para_latex(para):
    """
    Quét một đoạn văn, thay $...$ / $$...$$ bằng OMML thật.
    """
    full_text = "".join(r.text or "" for r in para.runs)
    if "$" not in full_text:
        return

    matches = list(_LATEX_PAT.finditer(full_text))
    if not matches:
        return

    # Tạo danh sách segments
    segments: list[tuple[str, str]] = []
    prev = 0
    for m in matches:
        if m.start() > prev:
            segments.append(("text", full_text[prev:m.start()]))
        latex = (m.group(1) or m.group(2) or "").strip()
        if latex:
            segments.append(("math", latex))
        prev = m.end()
    if prev < len(full_text):
        segments.append(("text", full_text[prev:]))

    # Xóa tất cả w:r khỏi đoạn (giữ w:pPr, bookmarks)
    p_elem = para._p
    keep_tags = {qn("w:pPr"), qn("w:bookmarkStart"), qn("w:bookmarkEnd")}
    for child in list(p_elem):
        if child.tag not in keep_tags:
            p_elem.remove(child)

    # Chèn lại nội dung mới
    for seg_type, content in segments:
        if seg_type == "text" and content:
            r    = OxmlElement("w:r")
            rPr  = OxmlElement("w:rPr")
            rF   = OxmlElement("w:rFonts")
            for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rF.set(qn(a), "Times New Roman")
            rPr.append(rF)
            for tag in ("w:sz", "w:szCs"):
                el = OxmlElement(tag); el.set(qn("w:val"), "28"); rPr.append(el)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = content
            r.append(t)
            p_elem.append(r)
        elif seg_type == "math" and content:
            omath = _latex_to_omath(content)
            if omath is not None:
                p_elem.append(omath)
            else:
                # Fallback: giữ nguyên text
                r = OxmlElement("w:r"); t = OxmlElement("w:t")
                t.text = f"${content}$"; r.append(t); p_elem.append(r)


def convert_latex_in_doc(doc: Document):
    """
    Duyệt toàn bộ tài liệu (thân bài + ô bảng), chuyển $...$ → OMML.
    Mỗi đoạn văn được bọc try/except riêng để lỗi không lan rộng.
    """
    import traceback

    def _safe(para):
        try:
            _process_para_latex(para)
        except Exception:
            logger.warning(f"latex para skip:\n{traceback.format_exc()}")

    for para in doc.paragraphs:
        _safe(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _safe(para)


# ═══════════════════════════════════════════════════════════════════════════════
# PHÂN TÍCH TIẾN TRÌNH BÀI DẠY — ENGINE
# ═══════════════════════════════════════════════════════════════════════════════

def _norm(text: str) -> str:
    """Bỏ dấu + lowercase để so sánh không phân biệt dấu."""
    nfkd = unicodedata.normalize("NFKD", text.lower())
    return "".join(c for c in nfkd if not unicodedata.combining(c))


# ── Nhận dạng loại hoạt động ──────────────────────────────────────────────────
# Mỗi loại có bộ từ khóa nhận dạng (không dấu, lowercase)
ACTIVITY_MARKERS: dict[str, list[str]] = {
    "khoi_dong": [
        "khoi dong", "mo dau", "dan nhap", "warm up", "tao hung thu",
        "gioi thieu bai", "vao bai", "kiem tra bai cu",
    ],
    "hinh_thanh": [
        "hinh thanh kien thuc", "tim hieu bai", "bai moi", "kien thuc moi",
        "noi dung bai", "nghien cuu", "kham pha", "tim hieu",
        "hoat dong 2", "hoat dong b",
    ],
    "luyen_tap": [
        "luyen tap", "cu co", "thuc hanh", "bai tap", "van dung ngay",
        "hoat dong 3", "hoat dong c",
    ],
    "van_dung": [
        "van dung", "mo rong", "lien he thuc te", "ung dung",
        "hoat dong 4", "hoat dong d",
    ],
    "danh_gia": [
        "danh gia", "tong ket", "nhan xet", "phan hoi", "kiem tra",
        "hoat dong 5",
    ],
}

ACTIVITY_NAMES: dict[str, str] = {
    "khoi_dong":  "Hoạt động Khởi động",
    "hinh_thanh": "Hoạt động Hình thành kiến thức",
    "luyen_tap":  "Hoạt động Luyện tập",
    "van_dung":   "Hoạt động Vận dụng",
    "danh_gia":   "Hoạt động Đánh giá / Tổng kết",
}

# Thứ tự hiển thị chuẩn của một tiến trình bài dạy
ACTIVITY_ORDER = ["khoi_dong", "hinh_thanh", "luyen_tap", "van_dung", "danh_gia"]


# ── Từ khóa nội dung → công cụ số gợi ý ─────────────────────────────────────
# key: từ khóa (không dấu), value: tên công cụ / gợi ý tiếng Việt đầy đủ
KEYWORD_TOOLS: dict[str, str] = {
    "thi nghiem":    "phần mềm mô phỏng thí nghiệm (PhET Interactive Simulations)",
    "bieu do":       "bảng tính (Excel/Google Sheets) để vẽ biểu đồ",
    "do thi":        "GeoGebra hoặc Desmos để vẽ và phân tích đồ thị",
    "tim hieu":      "tài nguyên số đáng tin cậy (thư viện số, Wikipedia khoa học)",
    "trinh bay":     "công cụ trình chiếu số (Google Slides / Canva / PowerPoint)",
    "thuyet trinh":  "phần mềm trình chiếu số (Google Slides / Canva)",
    "thao luan":     "nền tảng cộng tác số (Padlet, Google Docs, Jamboard)",
    "hop tac":       "công cụ cộng tác số (Google Workspace / Microsoft Teams)",
    "doc van ban":   "tài liệu số, e-book và tài nguyên số hóa",
    "viet":          "công cụ soạn thảo số (Google Docs / Word Online)",
    "tinh toan":     "phần mềm toán học số (GeoGebra / Desmos / Wolfram Alpha)",
    "tinh":          "máy tính số và phần mềm toán học (GeoGebra / Desmos)",
    "ve":            "phần mềm vẽ và thiết kế số (GeoGebra / Canva)",
    "quan sat":      "video số, hình ảnh vệ tinh và tài liệu đa phương tiện",
    "ban do":        "bản đồ số tương tác (Google Maps / Google Earth / GIS)",
    "mo phong":      "phần mềm mô phỏng (PhET / GeoGebra / Algodoo)",
    "kiem tra":      "công cụ kiểm tra số tương tác (Google Forms / Kahoot / Quizizz)",
    "trac nghiem":   "ứng dụng trắc nghiệm số (Quizizz / Kahoot / Quizlet)",
    "sang tao":      "công cụ sáng tạo số (Canva / Adobe Express)",
    "nghien cuu":    "cơ sở dữ liệu học thuật số (Google Scholar / thư viện số)",
    "hinh anh":      "kho hình ảnh số và phần mềm trực quan hóa",
    "video":         "video giáo dục số (YouTube Education / Khan Academy)",
    "phan tich":     "bảng tính và phần mềm phân tích dữ liệu số",
    "thu thap":      "công cụ thu thập dữ liệu số (Google Forms / khảo sát trực tuyến)",
    "bao cao":       "công cụ tạo báo cáo số (Google Docs / Canva)",
    "do luong":      "ứng dụng đo lường số và cảm biến kỹ thuật số",
    "tra cuu":       "từ điển số và cơ sở dữ liệu trực tuyến đáng tin cậy",
    "lap bang":      "bảng tính số (Excel / Google Sheets)",
    "so sanh":       "công cụ trực quan hóa số (biểu đồ, sơ đồ tư duy số)",
    "so do tu duy":  "phần mềm sơ đồ tư duy số (MindMeister / Canva / XMind)",
    "giai toan":     "phần mềm giải toán số (GeoGebra / Wolfram Alpha)",
    "bai tap":       "nền tảng luyện tập trực tuyến (Quizlet / Google Classroom)",
    "nhom":          "công cụ cộng tác nhóm số (Google Docs / Padlet / Miro)",
    "tong ket":      "sơ đồ tư duy số và công cụ trình bày tổng kết",
    "phan hoi":      "biểu mẫu phản hồi số (Google Forms / Mentimeter)",
}


# ── NLS theo loại hoạt động (template với placeholder {tool}) ────────────────
ACTIVITY_NLS_TEMPLATES: dict[str, list[str]] = {
    "khoi_dong": [
        "Khai thác video, hình ảnh số hoặc tình huống thực tiễn trên Internet để đặt vấn đề, tạo hứng thú và kết nối kiến thức đầu bài.",
        "Sử dụng công cụ khởi động tương tác (Kahoot / Mentimeter / Quizlet Live) để kiểm tra nhanh kiến thức nền và tạo không khí học tập sôi nổi.",
    ],
    "hinh_thanh": [
        "Dùng {tool} để minh họa, khám phá và kiểm chứng kiến thức mới một cách trực quan trong quá trình hình thành bài.",
        "Khai thác học liệu số đáng tin cậy (video bài giảng, animation, bài đọc số) để bổ sung và làm phong phú nội dung kiến thức.",
    ],
    "luyen_tap": [
        "Sử dụng {tool} để học sinh luyện tập theo nhịp độ cá nhân, nhận phản hồi tức thì và củng cố kiến thức hiệu quả.",
        "Khai thác nền tảng bài tập trực tuyến (Google Classroom / Quizlet) để giao và thu bài số, theo dõi tiến độ từng học sinh.",
    ],
    "van_dung": [
        "Tìm kiếm thông tin số về ứng dụng thực tiễn của kiến thức bài học; trình bày và chia sẻ kết quả bằng {tool}.",
        "Sử dụng công cụ số để đề xuất giải pháp cho tình huống thực tế — khuyến khích học sinh kết hợp kiến thức bài học với dữ liệu số.",
    ],
    "danh_gia": [
        "Dùng {tool} để kiểm tra nhanh, thu thập phản hồi học sinh và đánh giá hiệu quả bài dạy theo thời gian thực.",
        "Hướng dẫn học sinh tự đánh giá qua rubric số; giáo viên tổng kết bài bằng sơ đồ tư duy số hoặc slide.",
    ],
}

# Từ khóa ưu tiên tra cứu tool cho từng loại hoạt động
ACTIVITY_TOOL_PRIORITY: dict[str, list[str]] = {
    "hinh_thanh": ["thi nghiem", "mo phong", "do thi", "ban do", "ve", "tinh toan", "video"],
    "luyen_tap":  ["bai tap", "trac nghiem", "kiem tra", "tinh", "lap bang"],
    "van_dung":   ["trinh bay", "thuyet trinh", "bao cao", "sang tao", "nghien cuu"],
    "danh_gia":   ["kiem tra", "trac nghiem", "phan hoi"],
}


# ── Hàm phân tích chính ───────────────────────────────────────────────────────

def extract_doc_text(doc: Document) -> str:
    """Ghép toàn bộ text của tài liệu (thân bài + ô bảng)."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def detect_activities(text_norm: str) -> list[str]:
    """
    Trả về danh sách các loại hoạt động TÌM THẤY trong văn bản,
    theo đúng thứ tự ACTIVITY_ORDER.
    """
    found = []
    for act_type in ACTIVITY_ORDER:
        if any(marker in text_norm for marker in ACTIVITY_MARKERS[act_type]):
            found.append(act_type)
    return found


def find_keywords(text_norm: str) -> dict[str, str]:
    """Trả về {keyword: tool_description} cho các từ khóa tìm thấy."""
    return {kw: tool for kw, tool in KEYWORD_TOOLS.items() if kw in text_norm}


def build_activity_nls(act_type: str, found_kws: dict[str, str]) -> list[str]:
    """
    Tạo danh sách gợi ý NLS cho một hoạt động.
    Điền {tool} từ từ khóa nội dung phù hợp nhất.
    """
    templates = ACTIVITY_NLS_TEMPLATES.get(act_type, [])
    priorities = ACTIVITY_TOOL_PRIORITY.get(act_type, [])

    # Tìm tool phù hợp nhất theo thứ tự ưu tiên
    best_tool: str | None = None
    for kw in priorities:
        if kw in found_kws:
            best_tool = found_kws[kw]
            break
    if best_tool is None and found_kws:
        best_tool = next(iter(found_kws.values()))

    results = []
    for tpl in templates:
        if "{tool}" in tpl:
            tool = best_tool or "công cụ số phù hợp"
            results.append(tpl.replace("{tool}", tool))
        else:
            results.append(tpl)
    return results


def analyze_lesson_for_nls(doc: Document, mon: str, cap: str) -> dict:
    """
    Phân tích tài liệu, phát hiện hoạt động và gợi ý NLS từng hoạt động.

    Trả về:
    {
        "activities": [{"type": str, "name": str, "nls": [str, ...]}, ...],
        "keywords":   [str, ...],
        "fallback":   [str, ...]   # dùng khi không detect được activity nào
    }
    """
    text      = extract_doc_text(doc)
    text_norm = _norm(text)

    found_types = detect_activities(text_norm)
    found_kws   = find_keywords(text_norm)

    activities = []
    for act_type in found_types:
        activities.append({
            "type": act_type,
            "name": ACTIVITY_NAMES[act_type],
            "nls":  build_activity_nls(act_type, found_kws),
        })

    fallback: list[str] = []
    if not activities:
        # Không nhận dạng được tiến trình → dùng NLS chung theo môn/cấp
        fallback = NL_SO_DB.get(mon, {}).get(cap, [])

    return {
        "activities": activities,
        "keywords":   list(found_kws.keys()),
        "fallback":   fallback,
    }


# ── Chèn NLS có cấu trúc theo hoạt động vào .docx ────────────────────────────

def insert_smart_nls(doc: Document, analysis: dict, mon_label: str, cap_label: str):
    """
    Chèn khối NLS theo tiến trình bài dạy (hoặc NLS chung nếu fallback).
    Bỏ qua nếu đã tồn tại mục 'năng lực số'.
    """
    for para in doc.paragraphs:
        if "năng lực số" in para.text.lower():
            logger.info("Đã có 'Năng lực số' — bỏ qua chèn trùng.")
            return

    ref_p = _find_insertion_wp(doc)
    if ref_p is None:
        return

    # Tiêu đề chính (bold)
    header_text = (
        f"- Năng lực số tích hợp theo tiến trình bài dạy ({mon_label} – {cap_label}):"
        if analysis["activities"]
        else f"- Năng lực số tích hợp ({mon_label} – {cap_label}):"
    )
    current = ref_p
    header = _make_wp(header_text, bold=True)
    current.addnext(header); current = header

    if analysis["activities"]:
        for act in analysis["activities"]:
            # Sub-header hoạt động (bold, thụt lề nhẹ)
            sub = _make_wp(f"▸ {act['name']}:", bold=True, indent_twips=200)
            current.addnext(sub); current = sub
            # Các gợi ý NLS
            for item in act["nls"]:
                row = _make_wp(f"+ {item}", indent_twips=480)
                current.addnext(row); current = row
    else:
        # Fallback: danh sách phẳng
        for item in analysis["fallback"]:
            row = _make_wp(f"+ {item}", indent_twips=360)
            current.addnext(row); current = row


# ═══════════════════════════════════════════════════════════════════════════════
# FASTAPI ROUTES
# ═══════════════════════════════════════════════════════════════════════════════

@app.get("/", response_class=HTMLResponse)
async def serve_ui():
    html_path = os.path.join(os.path.dirname(__file__), "index.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/config")
async def get_config():
    """Public config cho frontend — chỉ trả về anon key (an toàn để public)."""
    return JSONResponse({
        "supabase_url":      SUPABASE_URL,
        "supabase_anon_key": SUPABASE_ANON_KEY,
        "auth_enabled":      AUTH_ENABLED,
        "free_quota":        FREE_QUOTA,
    })


@app.get("/me")
async def get_me(user: dict = Depends(get_current_user)):
    """Thông tin user + quota tháng hiện tại (dùng schema mới: used_quota, is_pro)."""
    uid  = user["sub"]
    supa = _get_supa()

    used   = 0
    is_pro = False

    if supa and AUTH_ENABLED and uid != "dev-local":
        current_month = datetime.now(timezone.utc).strftime("%Y-%m")
        try:
            r = (
                supa.table("profiles")
                .select("used_quota, is_pro, quota_month")
                .eq("id", uid)
                .maybe_single()
                .execute()
            )
            if r.data:
                is_pro = r.data.get("is_pro", False)
                # Reset nếu sang tháng mới
                if r.data.get("quota_month") == current_month:
                    used = r.data.get("used_quota", 0)
        except Exception as e:
            logger.warning(f"/me quota fetch: {e}")

    return JSONResponse({
        "email":  user.get("email", ""),
        "is_pro": is_pro,
        "usage":  used,
        "quota":  None if is_pro else FREE_QUOTA,
    })


@app.get("/competencies")
async def get_competencies(mon: str, cap: str):
    """Trả về danh sách năng lực số theo môn và cấp học."""
    mon_data = NL_SO_DB.get(mon)
    if not mon_data:
        raise HTTPException(status_code=404, detail=f"Môn '{mon}' không có trong database.")
    items = mon_data.get(cap, [])
    if not items:
        raise HTTPException(status_code=404,
                            detail=f"Chưa có dữ liệu cho {MON_LABELS.get(mon, mon)} – {CAP_LABELS.get(cap, cap)}.")
    return JSONResponse({"mon": MON_LABELS.get(mon, mon), "cap": CAP_LABELS.get(cap, cap),
                         "items": items})


def _cleanup(path: str):
    try: os.unlink(path)
    except: pass


@app.post("/process")
async def process_file(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    user: dict = Depends(get_current_user),
):
    check_and_increment_quota(user["sub"])
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .docx")
    tmp_in = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_in_path = tmp_in.name
    tmp_in.write(await file.read()); tmp_in.close()
    tmp_out_path = tmp_in_path.replace(".docx", "_out.docx")
    try:
        process_docx(tmp_in_path, tmp_out_path)
    except Exception as exc:
        import traceback
        logger.error(f"process_docx FAILED:\n{traceback.format_exc()}")
        _cleanup(tmp_in_path)
        raise HTTPException(status_code=500, detail=f"Lỗi xử lý: {type(exc).__name__}: {exc}")
    background_tasks.add_task(_cleanup, tmp_in_path)
    background_tasks.add_task(_cleanup, tmp_out_path)
    stem = os.path.splitext(file.filename)[0]
    return FileResponse(tmp_out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{stem}_chuanhoa.docx")


@app.post("/insert-nls")
async def insert_nls(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    mon:  str = Form(...),
    cap:  str = Form(...),
    user: dict = Depends(get_current_user),
):
    """Chèn Năng lực số theo GDPT 2018 vào giáo án .docx rồi trả về file."""
    check_and_increment_quota(user["sub"])
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .docx")

    mon_data = NL_SO_DB.get(mon)
    if not mon_data:
        raise HTTPException(status_code=404, detail=f"Không tìm thấy môn '{mon}'.")
    items = mon_data.get(cap, [])
    if not items:
        raise HTTPException(status_code=404,
                            detail=f"Chưa có dữ liệu cho {MON_LABELS.get(mon,mon)} – {CAP_LABELS.get(cap,cap)}.")

    tmp_in = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_in_path = tmp_in.name
    tmp_in.write(await file.read()); tmp_in.close()
    tmp_out_path = tmp_in_path.replace(".docx", "_nls.docx")

    try:
        doc = Document(tmp_in_path)
        insert_nang_luc_so(doc, items, MON_LABELS.get(mon, mon), CAP_LABELS.get(cap, cap))
        doc.save(tmp_out_path)
    except Exception as exc:
        _cleanup(tmp_in_path)
        raise HTTPException(status_code=500, detail=f"Lỗi chèn Năng lực số: {exc}")

    background_tasks.add_task(_cleanup, tmp_in_path)
    background_tasks.add_task(_cleanup, tmp_out_path)
    stem = os.path.splitext(file.filename)[0]
    return FileResponse(tmp_out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{stem}_nanglucso.docx")


@app.post("/analyze-nls-flow")
async def analyze_nls_flow_endpoint(
    file: UploadFile = File(...),
    mon:  str = Form(...),
    cap:  str = Form(...),
):
    """
    Phân tích tiến trình bài dạy trong file .docx.
    Trả về JSON với NLS gợi ý từng hoạt động.
    """
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ .docx")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.write(await file.read()); tmp.close()
    try:
        doc      = Document(tmp.name)
        analysis = analyze_lesson_for_nls(doc, mon, cap)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Lỗi phân tích: {exc}")
    finally:
        _cleanup(tmp.name)

    return JSONResponse({
        "mon":        MON_LABELS.get(mon, mon),
        "cap":        CAP_LABELS.get(cap, cap),
        "activities": analysis["activities"],
        "keywords":   analysis["keywords"],
        "fallback":   analysis["fallback"],
    })


@app.post("/insert-nls-smart")
async def insert_nls_smart(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    mon:  str = Form(...),
    cap:  str = Form(...),
    user: dict = Depends(get_current_user),
):
    """Phân tích tiến trình + chèn NLS theo từng hoạt động vào .docx."""
    check_and_increment_quota(user["sub"])
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ .docx")

    tmp_in = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_in_path = tmp_in.name
    tmp_in.write(await file.read()); tmp_in.close()
    tmp_out_path = tmp_in_path.replace(".docx", "_nls.docx")

    try:
        doc      = Document(tmp_in_path)
        analysis = analyze_lesson_for_nls(doc, mon, cap)
        insert_smart_nls(doc, analysis, MON_LABELS.get(mon, mon), CAP_LABELS.get(cap, cap))
        doc.save(tmp_out_path)
    except Exception as exc:
        _cleanup(tmp_in_path)
        raise HTTPException(status_code=500, detail=f"Lỗi chèn NLS: {exc}")

    background_tasks.add_task(_cleanup, tmp_in_path)
    background_tasks.add_task(_cleanup, tmp_out_path)
    stem = os.path.splitext(file.filename)[0]
    return FileResponse(
        tmp_out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{stem}_nanglucso.docx",
    )


# ═══════════════════════════════════════════════════════════════════════════════
# INTERACTIVE NLS — Upload Framework + Review + Inject
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text_from_upload(file_path: str, filename: str) -> str:
    """Trích xuất văn bản từ PDF / DOCX / TXT."""
    ext = os.path.splitext(filename.lower())[1]
    try:
        if ext == ".docx":
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif ext in (".txt", ".md"):
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        elif ext == ".pdf":
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                return "\n".join(pages)
    except Exception as e:
        logger.warning(f"extract_text_from_upload: {e}")
    return ""


def _extract_coded_items(text: str) -> list:
    """
    Quét văn bản khung NLS để tìm CÁC MỤC CÓ MÃ ĐỊNH DANH.
    Hỗ trợ các định dạng phổ biến trong chương trình VN:
      2.4.TC2a    → section.sub.TypeCodeLetter
      NL1.2, NL3b → TypeNumber.Sub
      TC2a, KNS4  → TypeCodeLetter
      1.1, 2.3.4  → Số thứ tự thuần túy
    Trả về danh sách [{"code": str, "text": "MÃ – Nội dung", "checked": True}, ...]
    """
    import uuid, re

    CODE_RE = re.compile(
        r'(?m)^[ \t]*'
        r'('
        # 1.1TC1a, 2.4TC2a (KHÔNG có dấu chấm giữa số và chữ — định dạng bang-tra-cuu-nls)
        r'\d+\.\d+[A-ZĐẮẶẪẨẦ]{1,5}\d+[a-z]?'
        # 2.4.TC2a, 1.1.NL3b (CÓ dấu chấm giữa số và chữ)
        r'|\d+\.\d+\.[A-ZĐẮẶẪẨẦ]{1,5}\d+[a-z]?'
        # NL1.2.3a, TC2.4
        r'|[A-ZĐẮẶẪẨẦ]{2,5}\d+\.\d+\.?\d*[a-z]?'
        # TC2a, NL3, KNS4b
        r'|[A-ZĐẮẶẪẨẦ]{2,5}\d+[a-z]?'
        # 1.1, 2.3.4 (chỉ số, ít ưu tiên hơn)
        r'|\d+\.\d+\.?\d*[a-z]?'
        r')'
        r'[ \t]*[).\-–:\s]+[ \t]*'
        r'(.{10,300})',
        re.UNICODE
    )

    items = []
    seen: set = set()

    for m in CODE_RE.finditer(text):
        code    = m.group(1).strip()
        content = m.group(2).strip()
        key     = code.upper()

        # Bỏ qua mã trùng hoặc nội dung quá ngắn
        if key in seen or len(content) < 15:
            continue
        # Bỏ qua dòng trông như metadata (tiêu đề, số tiết...)
        if re.search(r'\d+\s*tiết|\d{4}[-–]\d{4}|tuần|học kỳ', content, re.I):
            continue

        seen.add(key)
        items.append({
            "id":      str(uuid.uuid4())[:8],
            "code":    code,
            "text":    f"{code} – {content}",
            "checked": True,
        })

    return items


def _text_to_items(text: str) -> list:
    """
    Chuyển văn bản khung NLS → danh sách checkbox item.
    Chỉ lấy các dòng trông giống năng lực/kỹ năng thực sự.
    Lọc bỏ: tiêu đề ALL CAPS, thời gian, số tiết, mã môn, metadata.
    """
    import uuid, re

    # Từ khóa gợi ý đây là dòng NLS
    NLS_KEYWORDS = (
        "sử dụng", "khai thác", "phân tích", "đánh giá", "tạo lập",
        "giao tiếp", "hợp tác", "tìm kiếm", "thu thập", "trình bày",
        "nghiên cứu", "thiết kế", "lập trình", "ứng dụng", "giải quyết",
        "nhận biết", "phân biệt", "bảo vệ", "an toàn", "số", "digital",
        "công cụ", "phần mềm", "ứng dụng", "internet", "mạng", "dữ liệu",
        "năng lực", "kỹ năng", "thái độ", "kiến thức", "hiểu",
    )
    # Pattern loại trừ: dòng all caps, số tiết, mã năm học, tiêu đề section
    SKIP_PATTERNS = (
        re.compile(r"^[A-ZÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚĂĐĨŨƠƯẠ\s\d:;.,/–\-]{10,}$"),  # all caps
        re.compile(r"\d+\s*tiết"),           # "140 tiết", "18 tuần"
        re.compile(r"\d{4}\s*[-–]\s*\d{4}"), # năm học 2024-2025
        re.compile(r"^(môn|lớp|khối|học kỳ|hk|tuần|năm học|trường|tên|ngày)[\s:]+", re.I),
        re.compile(r"^[IVX]+\.\s"),          # I. II. III. mục lớn
        re.compile(r"^\d+\.\s+[A-ZÀÁÂ]"),   # 1. Phân phối...
        re.compile(r"^(kế hoạch|phân phối|nhiệm vụ|tiến độ)", re.I),
    )

    items = []
    for raw in text.splitlines():
        line = raw.strip().lstrip("•-*+○●▸▪►⁃–—◆→").strip()
        if len(line) < 20 or len(line) > 300:
            continue
        # Bỏ qua dòng bị loại trừ
        if any(p.search(line) for p in SKIP_PATTERNS):
            continue
        # Ưu tiên dòng có từ khóa NLS
        line_lower = line.lower()
        if not any(kw in line_lower for kw in NLS_KEYWORDS):
            continue
        items.append({"id": str(uuid.uuid4())[:8], "text": line, "checked": True})

    return items[:30]


def _match_activity_type(heading: str) -> str:
    """Xác định loại hoạt động từ tiêu đề heading."""
    hn = _norm(heading)
    for act_type, markers in ACTIVITY_MARKERS.items():
        if any(m in hn for m in markers):
            return act_type
    return "other"


def _extract_activities_with_nls(doc: Document, mon: str = "") -> list[dict]:
    """
    Phát hiện từng hoạt động học tập trong giáo án, trích xuất nội dung,
    rồi gợi ý mã NLS phù hợp với BỐI CẢNH CỤ THỂ của mỗi hoạt động.

    Pattern nhận dạng heading hoạt động:
      "Hoạt động 1:", "HĐ 2:", "HOẠT ĐỘNG A", "Hoạt động Khởi động", ...
    """
    import uuid

    ACT_HEAD_RE = re.compile(
        r'(?i)(?:hoạt\s*động|h[.đd])\s*(?:\d+|[a-z]|[ivx]+)',
        re.UNICODE
    )

    activities: list[dict] = []
    current_type: str | None = None
    current_name: str = ""
    current_texts: list[str] = []

    def _flush():
        """Lưu hoạt động hiện tại nếu có."""
        if current_type is None or not current_texts:
            return
        content   = " ".join(current_texts)
        suggested = _suggest_codes_from_lesson(_norm(content), mon=mon)
        if suggested:
            activities.append({
                "type":  current_type,
                "name":  current_name or "Hoạt động",
                "codes": suggested,
            })

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Phát hiện heading hoạt động
        if ACT_HEAD_RE.search(text) and len(text) < 120:
            _flush()
            current_type  = _match_activity_type(text)
            current_name  = text[:80]
            current_texts = []
        elif current_type is not None:
            # Dừng nếu gặp phần lớn khác (Mục tiêu, Thiết bị, ...)
            tl = text.lower()
            if any(kw in tl for kw in
                   ["ii. thiết bị", "iii. tiến trình", "a. mục tiêu",
                    "i. mục tiêu", "iv.", "v. "]):
                _flush()
                current_type = None
                current_texts = []
            else:
                current_texts.append(text)

    _flush()  # Lưu activity cuối

    # Gán id cho từng item
    for act in activities:
        for item in act["codes"]:
            item.setdefault("id", str(uuid.uuid4())[:8])

    return activities


def _suggest_codes_from_lesson(
    text_norm: str,
    max_codes: int = 5,
    mon: str = "",
) -> list[dict]:
    """
    Gợi ý tối đa max_codes mã NLS từ BANG_TRA_CUU.
    - Dựa trên từ khóa trong nội dung hoạt động.
    - Ưu tiên mã phù hợp với đặc thù môn học (mon).
    TUYỆT ĐỐI chỉ dùng mã từ BANG_TRA_CUU_MAP.
    """
    import uuid

    # Cấu hình ưu tiên theo môn
    subj_cfg         = SUBJECT_NLS_PRIORITY.get(mon, {})
    priority_codes   = set(subj_cfg.get("priority_codes", []))
    priority_secs    = set(subj_cfg.get("priority_sections", []))
    kw_boost_map     = subj_cfg.get("keyword_boost", {})

    seen: set[str]           = set()
    scored: list[tuple[int, str]] = []

    # ── Bước 1: Khớp từ khóa chung ───────────────────────────────────────────
    for kw, codes in KEYWORD_TO_NLS_CODES.items():
        kw_norm = _norm(kw)
        kw_len  = len(kw_norm)

        if kw_len <= 3:  # word-boundary check cho từ ngắn
            if not re.search(r'(?<![a-z])' + re.escape(kw_norm) + r'(?![a-z])', text_norm):
                continue
        else:
            if kw_norm not in text_norm:
                continue

        base = kw_len  # từ khóa dài → ưu tiên cao hơn
        for code in codes:
            if code in BANG_TRA_CUU_MAP and code not in seen:
                seen.add(code)
                # Bonus điểm nếu code phù hợp với môn học
                subj_bonus = 0
                if code in priority_codes:
                    subj_bonus = 60          # Mã nằm trong danh sách ưu tiên môn
                elif any(code.startswith(s) for s in priority_secs):
                    subj_bonus = 30          # Mã cùng section ưu tiên
                scored.append((base + subj_bonus, code))

    # ── Bước 2: Bonus thêm từ keyword_boost của môn ──────────────────────────
    for kw, codes in kw_boost_map.items():
        if _norm(kw) in text_norm:
            for code in codes:
                if code in BANG_TRA_CUU_MAP and code not in seen:
                    seen.add(code)
                    scored.append((120, code))  # Rất ưu tiên — đặc thù môn + nội dung

    # ── Bước 3: Đối chiếu MÔ TẢ CHỈ BÁO (pass thứ 2) ─────────────────────────
    # So sánh từ vựng hoạt động với MÔ TẢ của từng mã trong BANG_TRA_CUU.
    # Mã PHẢI tồn tại trong BANG_TRA_CUU_MAP — tuyệt đối không bịa thêm.
    act_words = {w for w in text_norm.split() if len(w) > 4}
    for entry in BANG_TRA_CUU:
        code = entry["code"]
        if code in seen:
            continue                      # đã được đưa vào từ pass 1
        if code not in BANG_TRA_CUU_MAP:
            continue                      # bảo vệ kép: chỉ mã hợp lệ

        desc_norm  = _norm(entry["content"])
        desc_words = {w for w in desc_norm.split() if len(w) > 4}
        overlap    = len(act_words & desc_words)

        if overlap < 2:                   # cần ít nhất 2 từ khóa trùng khớp
            continue

        seen.add(code)
        subj_bonus = (
            60 if code in priority_codes else
            (30 if any(code.startswith(s) for s in priority_secs) else 0)
        )
        scored.append((float(overlap * 5) + subj_bonus, code))

    # ── Bước 4: Sắp xếp, lấy top và xây kết quả ─────────────────────────────
    scored.sort(key=lambda x: -x[0])
    top_codes = [code for _, code in scored[:max_codes]]

    result: list[dict] = []
    for code in top_codes:
        # Xác nhận lần cuối: mã PHẢI trong BANG_TRA_CUU_MAP
        if code not in BANG_TRA_CUU_MAP:
            logger.warning(f"_suggest: bỏ qua mã không hợp lệ '{code}'")
            continue

        entry   = BANG_TRA_CUU_MAP[code]
        tools   = CODE_TO_TOOLS.get(code, [])
        content = entry["content"]

        # Xây explanation ngắn gọn
        match_kws = [kw for kw in KEYWORD_TO_NLS_CODES
                     if code in KEYWORD_TO_NLS_CODES[kw]
                     and _norm(kw) in text_norm]
        if match_kws:
            why = f"Hoạt động chứa '{match_kws[0]}' → phù hợp chỉ báo này"
        else:
            common = act_words & {w for w in _norm(content).split() if len(w) > 4}
            why = f"Từ khóa chung: {', '.join(list(common)[:3])}" if common else "Ưu tiên theo môn học"

        result.append({
            "id":          str(uuid.uuid4())[:8],
            "code":        code,
            "text":        f"{code} – {content}",
            "description": content,         # mô tả gốc từ BANG_TRA_CUU
            "explanation": why,             # lý do chọn mã này
            "tools":       tools,
            "checked":     True,
            "section":     entry["section"],
            "category":    entry["category"],
        })
    return result


def _detect_subject_from_doc(doc: Document) -> str:
    """
    Thử phát hiện môn học từ nội dung file (dòng 'MÔN: ...' ở đầu trang).
    Trả về mã môn tương ứng với SUBJECT_NLS_PRIORITY, hoặc '' nếu không tìm thấy.
    """
    MON_DETECT: dict[str, str] = {
        "toán":                 "toan",
        "ngữ văn":              "ngu_van",
        "tiếng anh":            "tieng_anh",
        "vật lý": "vat_ly",    "vật lí": "vat_ly",
        "hóa học":              "hoa_hoc",
        "sinh học":             "sinh_hoc",
        "khoa học tự nhiên":    "sinh_hoc",
        "khtn":                 "sinh_hoc",
        "lịch sử":              "lich_su",
        "địa lý": "dia_ly",    "địa lí": "dia_ly",
        "gdcd":                 "gdcd",
        "tin học":              "tin_hoc",
        "công nghệ":            "cong_nghe",
        "tự nhiên và xã hội":   "tnxh",
        "gdtc":                 "gdtc",
        "âm nhạc":              "am_nhac",
        "mỹ thuật":             "my_thuat",
        "khoa học":             "khoa_hoc",
    }
    # Chỉ quét 25 đoạn đầu (phần header/mục tiêu của giáo án)
    for para in doc.paragraphs[:25]:
        tl = para.text.strip().lower()
        if not tl:
            continue
        for keyword, mon_code in MON_DETECT.items():
            if keyword in tl:
                logger.info(f"Phát hiện môn: '{keyword}' → '{mon_code}'")
                return mon_code
    return ""


def build_interactive_nls(doc: Document, mon: str, cap: str, framework_text: str = "") -> dict:
    """
    Phân tích tài liệu → trả về cấu trúc NLS có thể chọn/sửa.
    Luôn dùng mã từ BANG_TRA_CUU (bang-tra-cuu-nls-cua-hs.pdf).
    Không bao giờ tự bịa mã mới.
    mon: mã môn học (dùng để ưu tiên mã NLS theo đặc thù môn).
    cap: cấp học (hiện dùng để log; bảng NLS dùng cùng mã cho mọi cấp).
    """
    import uuid
    logger.info(f"build_interactive_nls: môn={mon}, cấp={cap}")

    # Nếu có framework upload → chỉ chấp nhận mã khớp với BANG_TRA_CUU_MAP
    if framework_text.strip():
        coded = _extract_coded_items(framework_text)
        # Lọc chỉ giữ mã có trong BANG_TRA_CUU
        validated = []
        for c in coded:
            raw_code = c["code"].upper().replace(" ", "")
            # Tìm khớp trong map (case-insensitive)
            matched = next(
                (k for k in BANG_TRA_CUU_MAP if k.upper() == raw_code), None
            )
            if matched:
                entry = BANG_TRA_CUU_MAP[matched]
                validated.append({
                    "id":      c["id"],
                    "code":    matched,
                    "text":    f"{matched} – {entry['content']}",
                    "checked": True,
                    "section": entry["section"],
                })
            else:
                logger.warning(f"Mã '{c['code']}' không có trong BANG_TRA_CUU — bỏ qua.")

        if validated:
            logger.info(f"Framework: {len(validated)}/{len(coded)} mã khớp BANG_TRA_CUU.")
            return {
                "source":     "framework_coded",
                "has_codes":  True,
                "activities": [],
                "fallback":   validated,
                "keywords":   [],
            }
        logger.info("Framework không chứa mã khớp BANG_TRA_CUU — dùng phân tích bài dạy.")

    # Phân tích từng hoạt động học tập và gợi ý mã NLS theo bối cảnh cụ thể
    import uuid as _uuid
    text_norm  = _norm(extract_doc_text(doc))
    found_kws  = list(find_keywords(text_norm).keys())

    # Trích xuất TỪNG hoạt động + gợi ý mã phù hợp với nội dung + môn học
    act_sections = _extract_activities_with_nls(doc, mon=mon)

    if act_sections:
        # Chuyển codes → items theo định dạng frontend
        activities = []
        for act in act_sections:
            items = [
                {
                    "id":      c.get("id", str(_uuid.uuid4())[:8]),
                    "code":    c["code"],
                    "text":    c["text"],
                    "checked": True,
                }
                for c in act["codes"]
            ]
            activities.append({
                "type":  act["type"],
                "name":  act["name"],
                "items": items,
            })
        return {
            "source":     "btc_database",
            "has_codes":  True,
            "activities": activities,
            "fallback":   [],
            "keywords":   found_kws,
        }

    # Không phát hiện được từng hoạt động → gợi ý chung từ toàn bộ bài
    suggested = _suggest_codes_from_lesson(text_norm)
    return {
        "source":     "btc_database",
        "has_codes":  True,
        "activities": [],
        "fallback":   suggested,
        "keywords":   found_kws,
    }


# ── Helpers: chèn NLS vào bảng tiến trình ────────────────────────────────────

def _heading_matches(para_text: str, act_name: str) -> bool:
    """Kiểm tra đoạn văn có phải heading của hoạt động cần tìm không."""
    pt = _norm(para_text.strip())
    an = _norm(act_name.strip())
    if not pt or not an:
        return False
    # Khớp nếu 25 ký tự đầu giống nhau hoặc một bên chứa bên kia
    return an[:25] in pt or pt[:25] in an or (len(an) > 5 and an in pt)


def _is_activity_table(table) -> bool:
    """Kiểm tra bảng có phải bảng Hoạt động GV/HS không."""
    if not table.rows:
        return False
    row0 = _norm(" ".join(c.text for c in table.rows[0].cells))
    return any(kw in row0 for kw in
               ["giao vien", "hoc sinh", "hoat dong", " gv", " hs"])


def _add_nls_row_to_table(table, nls_items: list) -> None:
    """
    Thêm hàng 'Năng lực số tích hợp' vào cuối bảng hoạt động.
    Cột 1 (GV): nhãn in đậm nghiêng.
    Cột 2 (HS): danh sách mã NLS (tối đa 5 mã).
    """
    if not nls_items:
        return
    try:
        col_count = len(table.columns)
        row = table.add_row()
        cells = row.cells

        # ── Cột 1: Nhãn ─────────────────────────────────────────────────────
        p0 = cells[0].paragraphs[0]
        p0.clear()
        r0 = p0.add_run("Năng lực số tích hợp (NLS):")
        r0.bold   = True
        r0.italic = True
        r0.font.name = "Times New Roman"
        r0.font.size = Pt(13)

        # ── Cột 2: Mã NLS + Công cụ kết hợp ────────────────────────────────
        if col_count >= 2:
            c1    = cells[1]
            para  = c1.paragraphs[0]
            para.clear()
            first = True

            # Lấy loại hoạt động từ item đầu tiên → tra bảng công cụ theo loại HĐ
            act_type       = (nls_items[0].get("activity_type") or "other") if nls_items else "other"
            act_type_tools = ACTIVITY_TYPE_TOOLS.get(act_type, ACTIVITY_TYPE_TOOLS["other"])

            for item in nls_items[:5]:
                text     = (item.get("text") or "").strip()
                code     = (item.get("code") or "").strip()
                if not text:
                    continue

                # Công cụ: kết hợp (1) từ mã NLS + (2) từ loại HĐ, bỏ trùng
                code_tools = item.get("tools") or CODE_TO_TOOLS.get(code, [])
                extra      = [t for t in act_type_tools if t not in code_tools]
                all_tools  = code_tools + extra[:2]   # tối đa = 3 từ mã + 2 từ loại HĐ

                p = para if first else c1.add_paragraph()
                if first:
                    p.clear()
                first = False

                # Dòng 1: mã + nội dung năng lực (Times 13)
                r_nl = p.add_run(text)
                r_nl.font.name = "Times New Roman"
                r_nl.font.size = Pt(13)

                # Dòng 2: công cụ gợi ý cụ thể (in nghiêng, Times 12)
                if all_tools:
                    p_tool = c1.add_paragraph()
                    r_tool = p_tool.add_run(f"  ▶ Công cụ thực hiện: {'; '.join(all_tools[:4])}")
                    r_tool.italic    = True
                    r_tool.font.name = "Times New Roman"
                    r_tool.font.size = Pt(12)

        # Nếu có > 2 cột → merge cột 2 trở đi
        if col_count > 2:
            try:
                cells[1].merge(cells[col_count - 1])
            except Exception:
                pass

    except Exception as e:
        logger.warning(f"_add_nls_row_to_table: {e}")


def insert_nls_into_activity_tables(doc: Document, selected_items: list) -> None:
    """
    Duyệt body theo thứ tự:
      paragraph (heading hoạt động) → table kề tiếp → thêm hàng NLS.
    Chỉ xử lý các item CÓ activity_name.
    """
    from collections import defaultdict

    # Nhóm NLS theo tên hoạt động
    act_nls: dict[str, list] = defaultdict(list)
    for item in selected_items:
        act_name = (item.get("activity_name") or "").strip()
        if act_name:
            act_nls[act_name].append(item)

    if not act_nls:
        return

    body             = doc.element.body
    current_nls      = None
    current_act_key  = None
    processed: set   = set()

    for elem in body:
        # ── Đoạn văn ─────────────────────────────────────────────────────────
        if elem.tag == f"{{{W_NS}}}p":
            para_text = "".join(
                nd.text or "" for nd in elem.iter(f"{{{W_NS}}}t")
            ).strip()
            if not para_text:
                continue

            # Tìm hoạt động khớp
            matched_key = None
            for act_name in act_nls:
                if act_name not in processed and _heading_matches(para_text, act_name):
                    matched_key = act_name
                    break
            if matched_key:
                current_nls    = act_nls[matched_key]
                current_act_key = matched_key

        # ── Bảng ─────────────────────────────────────────────────────────────
        elif elem.tag == f"{{{W_NS}}}tbl" and current_nls:
            # Tìm Table object tương ứng
            tbl_obj = next((t for t in doc.tables if t._tbl is elem), None)

            if tbl_obj and _is_activity_table(tbl_obj):
                _add_nls_row_to_table(tbl_obj, current_nls)
                if current_act_key:
                    processed.add(current_act_key)
                current_nls    = None
                current_act_key = None


# ── Từ khóa NGHIÊM NGẶT: công cụ/hành động số cụ thể ───────────────────────
# (chặt hơn _DIGITAL_KW cũ — tránh false-positive như 'nhóm', 'sgk trang ...')
_DIGITAL_KW = [
    # Công cụ số cụ thể
    "video", "clip", "hình ảnh số",
    "phần mềm", "ứng dụng", "app",
    "máy tính", "điện thoại thông minh", "tablet",
    "internet", "trực tuyến", "online", "mạng internet",
    "slide", "powerpoint", "google slides", "canva",
    "google docs", "google forms", "padlet", "jamboard",
    "kahoot", "quizlet", "quizizz", "mentimeter",
    "zoom", "meet", "teams", "bảng tính excel",
    "sgk điện tử", "sách giáo khoa điện tử",
    "mô phỏng", "phét", "geogebra", "desmos",
    # Hành động số đặc trưng
    "tìm kiếm thông tin trên internet",
    "tra cứu trên mạng",
    "chia sẻ trực tuyến",
    "trình bày bằng slide",
    "thảo luận trực tuyến",
    "nộp bài online",
    "xử lý số liệu",
    "thu thập dữ liệu số",
]


def _group_by_buoc(paragraphs: list) -> list[dict]:
    """
    Nhóm các đoạn văn trong ô theo từng 'Bước' (hoặc 'Hoạt động').
    Trả về: [{"heading": str, "head_idx": int, "last_idx": int, "full_text": str}, ...]

    Nhận dạng heading: "Bước 1:", "Bước 2:", "B1:", "Hoạt động 1:", v.v.
    """
    STEP_RE = re.compile(
        r'^(?:bước|buoc|b\.?\s*|hoat dong|hoạt động)\s*\d+',
        re.IGNORECASE | re.UNICODE
    )

    groups:  list[dict] = []
    cur_head = None
    cur_head_idx = -1
    cur_texts: list[str] = []
    cur_last_idx = -1

    def _flush():
        if cur_head is not None:
            groups.append({
                "heading":   cur_head,
                "head_idx":  cur_head_idx,
                "last_idx":  cur_last_idx if cur_last_idx >= 0 else cur_head_idx,
                "full_text": " ".join(cur_texts),
            })

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue

        if STEP_RE.match(_norm(text)):
            _flush()
            cur_head      = text
            cur_head_idx  = i
            cur_texts     = [text]
            cur_last_idx  = i
        elif cur_head is not None:
            cur_texts.append(text)
            cur_last_idx = i

    _flush()
    return groups


def _is_to_chuc_table(table) -> bool:
    """Kiểm tra đây có phải bảng 'Tổ chức thực hiện' (có cột HĐ GV/HS)."""
    if not table.rows or len(table.rows) < 2:
        return False
    header = _norm(" ".join(c.text for c in table.rows[0].cells))
    return any(kw in header for kw in [
        "gv va hs", "giao vien", "hoc sinh",
        "hoat dong", "to chuc", "thuc hien",
    ])


def _find_gv_col(table) -> int:
    """Tìm chỉ số cột 'HĐ CỦA GV VÀ HS' (mặc định cột 0)."""
    if not table.rows:
        return 0
    for i, cell in enumerate(table.rows[0].cells):
        ct = _norm(cell.text)
        if any(kw in ct for kw in [
            "hd cua gv va hs", "gv va hs", "hoat dong cua gv",
            "to chuc thuc hien", "cua gv va hs",
        ]):
            return i
    return 0


def _cell_font_info(cell) -> tuple[str, float]:
    """Đọc font và cỡ chữ từ nội dung ô để đồng bộ định dạng."""
    for para in cell.paragraphs:
        for run in para.runs:
            if run.font.name and run.font.size:
                return run.font.name, run.font.size.pt
    return "Times New Roman", 13.0


def _has_nls_below(paragraphs: list, idx: int, window: int = 2) -> bool:
    """Kiểm tra trong [window] đoạn ngay sau idx đã có NLS chưa."""
    for j in range(idx + 1, min(idx + 1 + window, len(paragraphs))):
        if "nang luc so" in _norm(paragraphs[j].text):
            return True
    return False


def _make_nls_para_elem(text: str, bold=False, italic=False,
                         font_name="Times New Roman", font_sz_pt=13.0):
    """Tạo <w:p> để chèn inline vào ô bảng."""
    p  = OxmlElement("w:p")
    r  = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rF = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rF.set(qn(a), font_name)
    rPr.append(rF)

    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), str(int(font_sz_pt * 2)))
        rPr.append(el)

    if bold:   rPr.append(OxmlElement("w:b"))
    if italic: rPr.append(OxmlElement("w:i"))
    r.append(rPr)

    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def _insert_nls_block_inline(ref_elem, items: list,
                              font_name: str, font_sz: float,
                              label: str = "Năng lực số tích hợp"):
    """
    Chèn khối NLS gồm nhiều đoạn ngay SAU ref_elem (trong cùng ô bảng).
    Dùng addnext theo thứ tự ngược để thứ tự cuối cùng đúng.
    Kết quả: Label → Code1 → Tool1 → Code2 → Tool2 ...
    """
    lines: list[tuple[str, bool, bool]] = [
        (f"{label} (NLS):", True, True)
    ]
    for item in items[:2]:
        code  = item.get("code", "")
        text  = (item.get("text") or "").strip()
        tools = item.get("tools") or CODE_TO_TOOLS.get(code, [])
        if text:
            lines.append((text, False, False))
        if tools:
            lines.append((
                f"  ▶ Công cụ thực hiện: {'; '.join(tools[:2])}",
                False, True,
            ))

    # Chèn ngược để thứ tự cuối là đúng
    current = ref_elem
    for txt, bold, italic in reversed(lines):
        new_p = _make_nls_para_elem(txt, bold=bold, italic=italic,
                                     font_name=font_name, font_sz_pt=font_sz)
        current.addnext(new_p)


def _filter_items_for_para(text_norm: str, selected_items: list) -> list:
    """Lọc NLS items phù hợp nhất với nội dung đoạn văn."""
    scored = []
    for item in selected_items:
        item_norm = _norm(item.get("text", ""))
        score = sum(1 for w in item_norm.split() if len(w) > 3 and w in text_norm)
        if score > 0:
            scored.append((score, item))
    scored.sort(key=lambda x: -x[0])
    return [it for _, it in scored[:2]]


def _insert_nls_into_muc_tieu(
    doc: Document, selected_items: list, mon_label: str, cap_label: str
) -> None:
    """
    Chèn danh sách NLS tổng hợp vào cuối phần 'Năng lực' trong I. Mục tiêu.
    Bỏ qua nếu NLS đã tồn tại trong các đoạn văn thông thường (ngoài bảng).
    """
    # Kiểm tra xem đã có NLS ở phần văn bản thường chưa (không đếm bảng)
    for para in doc.paragraphs:
        if "năng lực số" in para.text.lower():
            logger.info("Mục tiêu đã có 'Năng lực số' — bỏ qua bước 1.")
            return

    ref_p = _find_insertion_wp(doc)
    if ref_p is None:
        return

    # Nhóm items theo activity_name để hiển thị có cấu trúc
    from collections import OrderedDict
    groups: OrderedDict = OrderedDict()
    ungrouped: list = []
    for item in selected_items:
        act  = (item.get("activity_name") or "").strip()
        text = (item.get("text") or "").strip()
        if not text:
            continue
        if act:
            groups.setdefault(act, []).append(text)
        else:
            ungrouped.append(text)

    # Chèn tiêu đề chính — dùng biến động subject_name (= mon_label đã resolve)
    if mon_label and cap_label:
        lbl = f"- Năng lực số tích hợp ({mon_label} – {cap_label}):"
    elif mon_label:
        lbl = f"- Năng lực số tích hợp ({mon_label}):"
    else:
        lbl = "- Năng lực số tích hợp:"
    header  = _make_wp(lbl, bold=True)
    current = ref_p
    current.addnext(header)
    current = header

    # Chèn theo nhóm hoạt động
    for act_name, texts in groups.items():
        sub = _make_wp(f"▸ {act_name}:", bold=True, indent_twips=200)
        current.addnext(sub); current = sub
        for text in texts:
            row = _make_wp(f"+ {text}", indent_twips=480)
            current.addnext(row); current = row

    # Chèn các mục không thuộc nhóm nào
    for text in ungrouped:
        row = _make_wp(f"+ {text}", indent_twips=360)
        current.addnext(row); current = row

    logger.info(f"Đã chèn NLS vào phần Năng lực (Mục tiêu): "
                f"{len(groups)} nhóm HĐ + {len(ungrouped)} mục chung.")


def inject_competence_to_docx(
    doc: Document, selected_items: list, mon_label: str, cap_label: str
):
    """
    Logic mới: Chèn NLS inline ngay sau từng bước/đoạn hoạt động
    có dùng công cụ số trong bảng 'Tổ chức thực hiện'.

    Nguyên tắc:
    1. Tìm bảng có cột 'HĐ CỦA GV VÀ HS'.
    2. Duyệt từng đoạn trong cột đó.
    3. Đoạn nào chứa từ khóa công cụ số → chèn NLS ngay bên dưới.
    4. Định dạng đồng bộ với ô bảng hiện có.
    5. Bỏ qua nếu NLS đã tồn tại ngay dưới đoạn đó.
    """
    # Phát hiện môn từ file (để dùng khi mon_label không được truyền vào)
    mon_code = _detect_subject_from_doc(doc)

    # Xây tên môn hiển thị: ưu tiên mon_label → phát hiện từ file → để trống
    subject_name = mon_label or MON_LABELS.get(mon_code, "")

    # ── BƯỚC 1: Chèn NLS tổng hợp vào cuối phần Năng lực (I. Mục tiêu) ──────
    # Truyền subject_name (đã giải quyết) thay vì mon_label gốc
    _insert_nls_into_muc_tieu(doc, selected_items, subject_name, cap_label)

    # ── BƯỚC 2: Chèn NLS vào cuối từng BƯỚC trong bảng 'Tổ chức thực hiện' ───
    # Nhóm nội dung theo Bước → phân tích toàn bộ Bước → chèn MỘT lần ở cuối.
    found_table  = False
    seen_tc_ids: set[int] = set()

    # Nhãn NLS dùng biến động — phản ánh đúng môn học từng bài
    label = (
        f"Năng lực số tích hợp – {subject_name}"
        if subject_name
        else "Năng lực số tích hợp"
    )

    for table in doc.tables:
        if not _is_to_chuc_table(table):
            continue

        gv_col      = _find_gv_col(table)
        found_table = True

        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:                     # bỏ qua header
                continue
            if len(row.cells) <= gv_col:
                continue

            cell = row.cells[gv_col]
            tc_id = id(cell._tc)
            if tc_id in seen_tc_ids:             # merged cell
                continue
            seen_tc_ids.add(tc_id)

            font_name, font_sz = _cell_font_info(cell)
            paras = list(cell.paragraphs)        # snapshot trước khi chèn

            # Nhóm đoạn theo từng Bước
            buoc_groups = _group_by_buoc(paras)

            if buoc_groups:
                # ── Có cấu trúc Bước → chèn cuối mỗi Bước ───────────────
                for group in buoc_groups:
                    # Kiểm tra Bước này đã có NLS chưa
                    last_para = paras[group["last_idx"]]
                    buoc_paras = paras[group["head_idx"]: group["last_idx"] + 1]
                    already    = any("năng lực số" in p.text.lower() for p in buoc_paras)
                    if already:
                        continue

                    # Phân tích TOÀN BỘ nội dung Bước
                    full_norm = _norm(group["full_text"])
                    has_dig   = any(_norm(kw) in full_norm for kw in _DIGITAL_KW)

                    # Nếu không rõ công cụ số → vẫn chèn nếu selected_items có mã
                    nls_items: list[dict] = []
                    if selected_items:
                        nls_items = _filter_items_for_para(full_norm, selected_items)
                    if not nls_items and has_dig:
                        nls_items = _suggest_codes_from_lesson(
                            full_norm, max_codes=2, mon=mon_code
                        )
                    if not nls_items:
                        continue

                    _insert_nls_block_inline(
                        last_para._p, nls_items,
                        font_name=font_name, font_sz=font_sz, label=label,
                    )
                    logger.info(f"NLS chèn cuối '{group['heading'][:50]}'")

            else:
                # ── Không có Bước → chèn cuối toàn bộ ô (nếu có công cụ số) ──
                full_norm = _norm(" ".join(p.text for p in paras))
                if not any(_norm(kw) in full_norm for kw in _DIGITAL_KW):
                    continue
                already = any("năng lực số" in p.text.lower() for p in paras)
                if already:
                    continue

                nls_items = selected_items or _suggest_codes_from_lesson(
                    full_norm, max_codes=2, mon=mon_code
                )
                if nls_items and paras:
                    _insert_nls_block_inline(
                        paras[-1]._p, nls_items[:2],
                        font_name=font_name, font_sz=font_sz, label=label,
                    )

    if not found_table:
        logger.info("Không tìm thấy bảng 'Tổ chức thực hiện'.")


# ── JSON builder cho format nghiêm ngặt ──────────────────────────────────────

def _build_activity_nls_json(
    activity_name: str, activity_text: str, mon: str = ""
) -> dict:
    """
    Phân tích một hoạt động và trả về JSON chuẩn:
    {activity_name, selected_codes: [{code, description}], explanation}

    Chỉ dùng mã có trong BANG_TRA_CUU_MAP.
    Tuyệt đối không bịa mã mới.
    """
    text_norm = _norm(activity_text)
    items     = _suggest_codes_from_lesson(text_norm, max_codes=3, mon=mon)

    selected_codes = []
    reasons: list[str] = []

    for item in items:
        code = item["code"]
        # Xác nhận lại mã hợp lệ trong BANG_TRA_CUU
        if code not in BANG_TRA_CUU_MAP:
            continue
        selected_codes.append({
            "code":        code,
            "description": item["description"],
            "section":     item["section"],
            "category":    item["category"],
        })
        reasons.append(f"{code}: {item.get('explanation', '')}")

    mon_name = MON_LABELS.get(mon, mon) if mon else "môn học"
    if selected_codes:
        expl = (
            f"Với môn {mon_name}, đã phân tích nội dung hoạt động và đối chiếu "
            f"với bảng tra cứu NLS. Chọn {len(selected_codes)} mã phù hợp nhất: "
            + "; ".join(reasons[:3])
        )
    else:
        expl = (
            f"Hoạt động này không tích hợp được năng lực số cụ thể nào "
            f"từ bảng tra cứu bang-tra-cuu-nls-cua-hs.pdf."
        )

    return {
        "activity_name":  activity_name,
        "selected_codes": selected_codes,
        "explanation":    expl,
    }


# ── Endpoints ─────────────────────────────────────────────────────────────────

@app.post("/analyze-nls-strict")
async def analyze_nls_strict_endpoint(
    file: UploadFile = File(...),
    mon:  str = Form(""),
    cap:  str = Form(""),
    user: dict = Depends(get_current_user),
):
    """
    Phân tích giáo án và trả về mã NLS chính xác theo JSON format.

    - Chỉ dùng mã từ bang-tra-cuu-nls-cua-hs.pdf (109 mã).
    - Mỗi hoạt động trả về: {activity_name, selected_codes, explanation}.
    - selected_codes = [] nếu không tìm thấy mã phù hợp.
    """
    uid = user["sub"]
    logger.info(f"analyze-nls-strict: user={uid}, file={file.filename}, mon={mon}")
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Chỉ hỗ trợ .docx")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.write(await file.read()); tmp.close()
    results: list[dict] = []

    try:
        doc      = Document(tmp.name)
        detected = mon or _detect_subject_from_doc(doc)
        acts     = _extract_activities_with_nls(doc, mon=detected)

        if not acts:
            # Không phát hiện được hoạt động cụ thể → phân tích toàn bài
            full_text = extract_doc_text(doc)
            results.append(_build_activity_nls_json(
                "Toàn bộ bài dạy", full_text, mon=detected
            ))
        else:
            for act in acts:
                act_text = act.get("content", act.get("name", ""))
                result   = _build_activity_nls_json(
                    act["name"], act_text, mon=detected
                )
                results.append(result)

    except Exception as exc:
        raise HTTPException(500, f"Lỗi phân tích: {exc}")
    finally:
        _cleanup(tmp.name)

    return JSONResponse({
        "source":     "bang-tra-cuu-nls-cua-hs.pdf",
        "mon":        MON_LABELS.get(mon, mon),
        "cap":        CAP_LABELS.get(cap, cap),
        "activities": results,
        "total_codes_available": len(BANG_TRA_CUU),
    })


@app.get("/nls-codes")
async def get_nls_codes():
    """
    Trả về toàn bộ 109 mã NLS từ bang-tra-cuu-nls-cua-hs.pdf.
    Đây là nguồn duy nhất — frontend dùng để hiện dropdown picker.
    """
    return JSONResponse({
        "total":      len(BANG_TRA_CUU),
        "source":     "bang-tra-cuu-nls-cua-hs.pdf",
        "categories": CATEGORY_NAMES,
        "codes":      [
            {
                "code":      item["code"],
                "content":   item["content"],
                "full_text": f"{item['code']} – {item['content']}",
                "section":   item["section"],
                "category":  item["category"],
            }
            for item in BANG_TRA_CUU
        ],
    })


@app.post("/upload-framework")
async def upload_framework_endpoint(file: UploadFile = File(...)):
    """Upload file khung NLS (PDF/DOCX/TXT) và trả về văn bản đã trích xuất."""
    allowed = {".pdf", ".docx", ".txt", ".md"}
    ext = os.path.splitext(file.filename.lower())[1]
    if ext not in allowed:
        raise HTTPException(400, f"Chỉ hỗ trợ: {', '.join(sorted(allowed))}")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    tmp.write(await file.read())
    tmp.close()
    try:
        text = extract_text_from_upload(tmp.name, file.filename)
    finally:
        _cleanup(tmp.name)

    if not text.strip():
        raise HTTPException(422, "Không trích xuất được nội dung. Thử file khác nhé!")

    return JSONResponse({"text": text, "char_count": len(text)})


@app.post("/analyze-nls-interactive")
async def analyze_nls_interactive_endpoint(
    file: UploadFile = File(...),
    mon: str = Form(...),
    cap: str = Form(...),
    framework_text: str = Form(""),
    user: dict = Depends(get_current_user),
):
    """Phân tích bài dạy → trả về danh sách NLS dạng checkbox."""
    uid = user["sub"]   # dùng để log, không tính quota (phân tích = miễn phí)
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Chỉ hỗ trợ .docx")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.write(await file.read())
    tmp.close()
    try:
        doc    = Document(tmp.name)
        result = build_interactive_nls(doc, mon, cap, framework_text)
        logger.info(f"analyze-nls-interactive: user={uid}, source={result.get('source')}")
    except Exception as exc:
        raise HTTPException(500, f"Lỗi phân tích: {exc}")
    finally:
        _cleanup(tmp.name)

    return JSONResponse({
        "mon": MON_LABELS.get(mon, mon),
        "cap": CAP_LABELS.get(cap, cap),
        **result,
    })


@app.post("/inject-nls-selected")
async def inject_nls_selected_endpoint(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    mon:        str = Form(...),
    cap:        str = Form(...),
    items_json: str = Form(...),
    user: dict = Depends(get_current_user),
):
    """Nhận danh sách NLS đã duyệt → chèn vào .docx và trả về file."""
    import json as _json
    check_and_increment_quota(user["sub"])

    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Chỉ hỗ trợ .docx")

    try:
        selected = _json.loads(items_json)
    except Exception:
        raise HTTPException(400, "items_json không hợp lệ (phải là JSON)")

    if not selected:
        raise HTTPException(400, "Chưa chọn năng lực số nào")

    tmp_in = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_in_path = tmp_in.name
    tmp_in.write(await file.read())
    tmp_in.close()
    tmp_out = tmp_in_path.replace(".docx", "_nls.docx")

    try:
        doc = Document(tmp_in_path)
        inject_competence_to_docx(
            doc, selected,
            MON_LABELS.get(mon, mon),
            CAP_LABELS.get(cap, cap),
        )
        doc.save(tmp_out)
    except Exception as exc:
        _cleanup(tmp_in_path)
        raise HTTPException(500, f"Lỗi chèn NLS: {exc}")

    background_tasks.add_task(_cleanup, tmp_in_path)
    background_tasks.add_task(_cleanup, tmp_out)
    stem = os.path.splitext(file.filename)[0]
    return FileResponse(
        tmp_out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{stem}_nanglucso.docx",
    )


if __name__ == "__main__":
    # Local development only — production dùng gunicorn (xem Procfile)
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)
